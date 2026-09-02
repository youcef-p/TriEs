from flask import Flask, render_template, request, jsonify, send_file, Response
import os
import re
import io
import tempfile
import string
import platform
import time
import json
import copy
import queue
import threading
import shutil
import uuid
from pathlib import Path
from datetime import date, datetime
from typing import Any, Optional
import unicodedata

# Safely import external libraries
try: import pdfplumber
except ImportError: pdfplumber = None
try:
    import xlrd
    from openpyxl import load_workbook
    from openpyxl.styles import PatternFill
except ImportError:
    xlrd = None
    load_workbook = None

# Monthly reports imports
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
import logging
import logging.handlers
from dataclasses import dataclass
from difflib import SequenceMatcher


app = Flask(__name__)

# ==========================================
# HUB FUNCTIONS (original Nexus)
# ==========================================


# ==========================================
# APP 1: DRIVE LISTER LOGIC
# ==========================================
def format_size(size_bytes):
    for unit in ["B", "KB", "MB", "GB", "TB"]:
        if size_bytes < 1024: return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} PB"

def list_contents(path, show_hidden, show_sizes, recursive):
    lines = []
    path = Path(path).expanduser()
    if not path.exists() or not path.is_dir(): return [f"❌ '{path}' isn't a valid, accessible folder."]
    lines.append(f"\n📂 Listing contents of: {path}\n")
    lines.append("-" * 60)
    entry_count = 0; total_size = 0
    walker = os.walk(str(path)) if recursive else [(str(path), [], os.listdir(str(path)))]
    for root, dirs, files in walker:
        if not recursive:
            all_entries = os.listdir(str(path))
            dirs = [e for e in all_entries if (path / e).is_dir()]
            files = [e for e in all_entries if (path / e).is_file()]
        if not show_hidden:
            dirs[:] = [d for d in dirs if not d.startswith(".")]
            files = [f for f in files if not f.startswith(".")]
        rel_root = os.path.relpath(root, str(path))
        indent = "" if rel_root == "." else "  " * rel_root.count(os.sep)
        if rel_root != ".": lines.append(f"{indent}📁 {os.path.basename(root)}/")
        for d in sorted(dirs):
            if recursive: continue
            lines.append(f"{indent}📁 {d}/"); entry_count += 1
        for f in sorted(files):
            full_path = os.path.join(root, f)
            try: size = os.path.getsize(full_path)
            except OSError: size = 0
            total_size += size; entry_count += 1
            if show_sizes: lines.append(f"{indent}  📄 {f}  ({format_size(size)})")
            else: lines.append(f"{indent}  📄 {f}")
        if not recursive: break
    lines.append("-" * 60)
    lines.append(f"\n✅ Done! Found {entry_count} item(s).")
    if show_sizes: lines.append(f"📦 Total size of listed files: {format_size(total_size)}")
    return lines

def build_file_index(path, show_hidden):
    file_index = {}; dir_set = set()
    for root, dirs, files in os.walk(str(path)):
        if not show_hidden:
            dirs[:] = [d for d in dirs if not d.startswith(".")]
            files = [f for f in files if not f.startswith(".")]
        rel_root = os.path.relpath(root, str(path))
        for d in dirs: dir_set.add((d if rel_root == "." else os.path.join(rel_root, d)).replace(os.sep, "/"))
        for f in files:
            rel_file = (f if rel_root == "." else os.path.join(rel_root, f)).replace(os.sep, "/")
            try: size = os.path.getsize(os.path.join(root, f))
            except OSError: size = -1
            file_index[rel_file] = size
    return file_index, dir_set

def compare_listings(path_a, path_b, show_hidden):
    lines = []
    files_a, dirs_a = build_file_index(path_a, show_hidden)
    files_b, dirs_b = build_file_index(path_b, show_hidden)
    lines.append("             FOLDER / DRIVE COMPARISON RESULT\n" + "=" * 60 + f"\nA: {path_a}\nB: {path_b}\n" + "-" * 60)
    files_a_set, files_b_set = set(files_a.keys()), set(files_b.keys())
    only_in_a, only_in_b = sorted(files_a_set - files_b_set), sorted(files_b_set - files_a_set)
    common = sorted(files_a_set & files_b_set)
    same_size = [f for f in common if files_a[f] == files_b[f]]
    diff_size = [f for f in common if files_a[f] != files_b[f]]
    only_dirs_a, only_dirs_b = sorted(dirs_a - dirs_b), sorted(dirs_b - dirs_a)
    lines.append(f"\n📊 SUMMARY\n  Files only in A:            {len(only_in_a)}\n  Files only in B:            {len(only_in_b)}\n  Files in both, same size:   {len(same_size)}\n  Files in both, diff. size:  {len(diff_size)}\n  Folders only in A:          {len(only_dirs_a)}\n  Folders only in B:          {len(only_dirs_b)}")
    
    if only_dirs_a:
        lines.append(f"\n📁 Folders only in A ({len(only_dirs_a)}):")
        for d in only_dirs_a: lines.append(f"  + {d}/")
    if only_dirs_b:
        lines.append(f"\n📁 Folders only in B ({len(only_dirs_b)}):")
        for d in only_dirs_b: lines.append(f"  + {d}/")
    if only_in_a:
        lines.append(f"\n📄 Files only in A ({len(only_in_a)}):")
        for f in only_in_a: lines.append(f"  + {f}")
    if only_in_b:
        lines.append(f"\n📄 Files only in B ({len(only_in_b)}):")
        for f in only_in_b: lines.append(f"  + {f}")
    if diff_size:
        lines.append(f"\n⚠️  Different sizes ({len(diff_size)}):")
        for f in diff_size: lines.append(f"  ≠ {f}   A: {format_size(files_a[f])}   |   B: {format_size(files_b[f])}")
        
    lines.append(f"\n✅ Matches: {len(same_size)}\n" + "-" * 60)
    total_diffs = len(only_in_a) + len(only_in_b) + len(diff_size) + len(only_dirs_a) + len(only_dirs_b)
    lines.append("🎉 Perfect match." if total_diffs == 0 else f"🔎 Total differences: {total_diffs}")
    return lines

_HEADER_RE = re.compile(r"^📂 Listing contents of:\s*(.+)$")
_FOLDER_RE = re.compile(r"^(\s*)📁\s+(.+?)/\s*$")
_FILE_RE = re.compile(r"^(\s*)📄\s+(.+?)(?:\s{2}\(([^)]+)\))?\s*$")

def parse_listing_file(filepath):
    with open(filepath, "r", encoding="utf-8", errors="replace") as fh: raw_lines = fh.read().splitlines()
    source_label, files, dirs, stack, had_sizes = None, {}, set(), [], False
    for line in raw_lines:
        if source_label is None:
            m = _HEADER_RE.match(line.strip())
            if m: source_label = m.group(1).strip(); continue
        fm = _FOLDER_RE.match(line)
        if fm: depth = len(fm.group(1)) // 2; stack = stack[:depth] + [fm.group(2)]; dirs.add("/".join(stack)); continue
        pm = _FILE_RE.match(line)
        if pm:
            parent = "/".join(stack[:len(pm.group(1)) // 2])
            rel_path = f"{parent}/{pm.group(2)}" if parent else pm.group(2)
            if pm.group(3): had_sizes = True
            files[rel_path] = pm.group(3)
    return {"source_label": source_label or Path(filepath).name, "files": files, "dirs": dirs, "had_sizes": had_sizes}

def compare_parsed_listings(parsed_a, parsed_b):
    lines = [f"     COMPARISON OF TWO PREVIOUSLY EXPORTED LISTINGS\n{'=' * 60}\nA: {parsed_a['source_label']}\nB: {parsed_b['source_label']}\n{'-' * 60}"]
    both_sizes = parsed_a['had_sizes'] and parsed_b['had_sizes']
    files_a_set, files_b_set = set(parsed_a['files']), set(parsed_b['files'])
    only_a, only_b = sorted(files_a_set - files_b_set), sorted(files_b_set - files_a_set)
    common = sorted(files_a_set & files_b_set)
    same_size = [f for f in common if not both_sizes or parsed_a['files'][f] == parsed_b['files'][f]]
    diff_size = [f for f in common if both_sizes and parsed_a['files'][f] != parsed_b['files'][f]]
    lines.append(f"\n📊 SUMMARY\n  Files only in A: {len(only_a)}\n  Files only in B: {len(only_b)}")
    if only_a:
        lines.append(f"\n📄 Files only in A:")
        for f in only_a: lines.append(f"  + {f}")
    if only_b:
        lines.append(f"\n📄 Files only in B:")
        for f in only_b: lines.append(f"  + {f}")
    if diff_size:
        lines.append(f"\n⚠️  Different sizes:")
        for f in diff_size: lines.append(f"  ≠ {f}   A: {parsed_a['files'][f]}   |   B: {parsed_b['files'][f]}")
    lines.append(f"\n✅ Matches: {len(same_size)}\n{'-' * 60}")
    total_diffs = len(only_a) + len(only_b) + len(diff_size)
    lines.append("🎉 Perfect match." if total_diffs == 0 else f"🔎 Total differences: {total_diffs}")
    return lines

# ==========================================
# APP 3: PDF COMPRESSOR LOGIC
# ==========================================
STOPWORDS = {"the", "a", "an", "is", "are", "was", "were", "be", "been", "being", "have", "has", "had", "do", "does", "did", "will", "would", "could", "should", "may", "might", "must", "shall", "can", "need", "dare", "ought", "used", "to", "of", "in", "for", "on", "with", "at", "by", "from", "as", "into", "through", "during", "before", "after", "above", "below", "between", "under", "again", "further", "then", "once", "here", "there", "when", "where", "why", "how", "all", "each", "few", "more", "most", "other", "some", "such", "no", "nor", "not", "only", "own", "same", "so", "than", "too", "very", "just", "and", "but", "if", "or", "because", "until", "while", "about", "against", "between", "into", "through", "during", "before", "after", "above", "below", "up", "down", "out", "off", "over", "under", "again", "further", "then", "once"}
BOILERPLATE_PATTERNS = [r"page\s*\d+\s*of\s*\d+", r"\d+\s*/\s*\d+", r"confidential", r"all rights reserved", r"copyright\s*[©®™]?\s*\d{4}", r" proprietary ", r"terms of use", r"privacy policy", r"\bwww\.\S+\b", r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", r"tel[:\.]?\s*\+?[\d\s\-\(\)]{7,}", r"fax[:\.]?\s*\+?[\d\s\-\(\)]{7,}"]
URL_PATTERN = re.compile(r"https?://\S+|www\.\S+")
WHITESPACE_PATTERN = re.compile(r"\s+")

def normalize_text(text): return WHITESPACE_PATTERN.sub(" ", unicodedata.normalize("NFKC", text)).strip()
def remove_boilerplate(text):
    text = URL_PATTERN.sub("", text)
    for pattern in BOILERPLATE_PATTERNS: text = re.sub(pattern, "", text, flags=re.IGNORECASE)
    return text
def compress_paragraph(text, aggressive=False):
    words = text.split()
    filtered = [w for w in words if w.lower() not in STOPWORDS and len(w) > 1] if aggressive else [w for w in words if w.lower() not in STOPWORDS or len(w) > 3]
    return " ".join(filtered)
def table_to_markdown(table):
    if not table or len(table) < 1: return ""
    lines = []
    for i, row in enumerate(table):
        cells = [re.sub(r"\s+", " ", str(cell).strip() if cell is not None else "") for cell in row]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0: lines.append("|" + "|".join([" --- " for _ in cells]) + "|")
    return "\n".join(lines)
def extract_page_content(page, aggressive=False):
    chunks = []; tables = page.find_tables()
    text = page.extract_text()
    if text:
        text = remove_boilerplate(normalize_text(text))
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        compressed_paras = [compress_paragraph(p, aggressive) for p in paragraphs]
        chunks.append("\n\n".join([p for p in compressed_paras if p and len(p) > 3]))
    for table in tables:
        md_table = table_to_markdown(table.extract())
        if md_table: chunks.append("\n\n" + md_table + "\n")
    return "\n\n".join(chunks)

# ==========================================
# APP 2: SEISMIC DATA MINER LOGIC
# ==========================================
OUTPUT_COLUMNS = ("study_name", "reception_date", "support_type", "support_number", "swath_profile", "ffid", "vp", "data_type", "centre_traitement", "OBESERVATION", "source_file")
SUPPORT_COLOURS = {"Cassette": PatternFill("solid", fgColor="FFF2CC"), "CD/DVD": PatternFill("solid", fgColor="DEEAF6"), "USB": PatternFill("solid", fgColor="E2EFDA"), "DOC": PatternFill("solid", fgColor="F2F2F2"), "Unknown": PatternFill("solid", fgColor="FFFFFF")} if load_workbook else {}
FAMILY_SIT2, FAMILY_SIT1, FAMILY_STACK, FAMILY_RECEPTION, FAMILY_TEMPLATE = "SIT2", "SIT1", "STACK", "RECEPTION", "TEMPLATE"

class SeismicRecord:
    __slots__ = OUTPUT_COLUMNS
    def __init__(self):
        for col in OUTPUT_COLUMNS: setattr(self, col, None)

def make_record(study_name, reception_date, support_type, support_number, swath_profile="", data_type="", centre_traitement="", observation="", source_file=""):
    r = SeismicRecord()
    r.study_name = _strip(study_name) or None; r.reception_date = reception_date; r.support_type = support_type
    r.support_number = max(support_number, 1); r.swath_profile = _strip(swath_profile) or None; r.ffid = None; r.vp = None
    r.data_type = _strip(data_type) or None; r.centre_traitement = _strip(centre_traitement) or None
    r.OBESERVATION = _strip(observation) or None; r.source_file = source_file
    return r

def _emit_supports(n_cass, n_cd, n_usb, n_doc, study, rec_date, swath, data_type, centre, observation, fname):
    rows = []
    for i in range(n_cass): rows.append(make_record(study, rec_date, "Cassette", i + 1, swath, data_type, centre, observation, fname))
    for i in range(n_cd): rows.append(make_record(study, rec_date, "CD/DVD", i + 1, swath, data_type, centre, observation, fname))
    for i in range(n_usb): rows.append(make_record(study, rec_date, "USB", i + 1, swath, data_type, centre, observation, fname))
    for i in range(n_doc): rows.append(make_record(study, rec_date, "DOC", i + 1, swath, data_type, centre, observation, fname))
    return rows

def _strip(val):
    if val is None: return ""
    s = str(val).strip()
    return "" if s in ("None", "nan") else s

def _to_int(val):
    try: return int(float(val))
    except: return 0

def _xldate(serial, book):
    if not serial or not isinstance(serial, (int, float)): return None
    try:
        y, m, d, *_ = xlrd.xldate_as_tuple(serial, book.datemode)
        return date(y, m, d) if y else None
    except: return None

def _to_date(raw, book=None):
    if raw is None: return None
    if isinstance(raw, datetime): return raw.date()
    if isinstance(raw, date): return raw
    if isinstance(raw, (int, float)) and book is not None: return _xldate(raw, book)
    return None

def _normalise_support(raw):
    r = raw.upper()
    if any(k in r for k in ("CASS", "3592", "CASSETTE", "CART")): return "Cassette"
    if any(k in r for k in ("CD", "DVD", "DISC", "DISQUE")): return "CD/DVD"
    if any(k in r for k in ("USB", "CLÉ", "CLE", "FLASH")): return "USB"
    if "DOC" in r: return "DOC"
    return "Unknown"

def _join_obs(*parts): return " / ".join(p.strip() for p in parts if p and p.strip())

def _xls_obs(structure, trailing, sheet_name):
    parts = []
    s = _strip(structure)
    if s and len(s) > 1: parts.append(f"Structure : {s}")
    if sheet_name.strip().upper() == "ASS": parts.append("ASSOCIATION")
    t = _strip(trailing)
    if t: parts.append(t)
    return _join_obs(*parts)

class SeismicConsolidator:
    def __init__(self, log_cb): self.log_cb = log_cb

    def _is_template_xls(self, filepath: Path) -> bool:
        if filepath.suffix.lower() != ".xls": return False
        try:
            book = xlrd.open_workbook(str(filepath)); sheet = book.sheets()[0]
            for r in range(min(10, sheet.nrows)):
                if "directions" in " ".join(_strip(sheet.cell_value(r, c)).lower() for c in range(sheet.ncols)) and "date de rcep" in " ".join(_strip(sheet.cell_value(r, c)).lower() for c in range(sheet.ncols)): return True
        except: pass
        return False

    def _is_output_template(self, filepath: Path) -> bool:
        try:
            wb = load_workbook(str(filepath), read_only=True, data_only=True); val = wb.active.cell(1, 1).value; wb.close()
            return bool(val and _strip(val).lower() == "study_name")
        except: return False

    def detect_family(self, filepath: Path) -> Optional[str]:
        suffix = filepath.suffix.lower()
        if suffix == ".xlsx": return None if self._is_output_template(filepath) else FAMILY_RECEPTION
        if suffix == ".xls":
            if self._is_template_xls(filepath): return FAMILY_TEMPLATE
            stem = filepath.stem.lower()
            if "stack" in stem: return FAMILY_STACK
            if stem.startswith("sit1"): return FAMILY_SIT1
            return FAMILY_SIT2
        return None

    def discover_sources(self, directory: Path, output_path: Path) -> list:
        return [f for f in sorted(directory.glob("*.*")) if f.resolve() != output_path.resolve() and self.detect_family(f) is not None]

    def find_reception_template(self, directory: Path, output_path: Path) -> Path:
        candidates = [f for f in sorted(directory.glob("*.xlsx")) if f.resolve() != output_path.resolve() and self._is_output_template(f)]
        if not candidates: raise FileNotFoundError("No .xlsx output template found. Expected A1 == 'study_name'.")
        return candidates[0]

    def parse_sit2(self, filepath: Path) -> list:
        records = []; fname = filepath.name
        try: book = xlrd.open_workbook(str(filepath))
        except Exception as e: self.log_cb(f"⚠️ Cannot open {fname}: {e}"); return records
        for sheet_name in [s for s in ("SH", "ASS") if s in book.sheet_names()]:
            sheet = book.sheet_by_name(sheet_name); hdr_row = 0
            for r in range(min(10, sheet.nrows)):
                if any("campagne" in _strip(sheet.cell_value(r, c)).lower() for c in range(sheet.ncols)): hdr_row = r; break
            hdr = [_strip(sheet.cell_value(hdr_row, c)).lower() for c in range(sheet.ncols)]
            try: camp_col = next(i for i, h in enumerate(hdr) if "campagne" in h)
            except StopIteration: continue
            has_structure = camp_col > 0; trailing_start = camp_col + 6
            for row_idx in range(hdr_row + 1, sheet.nrows):
                row = [sheet.cell_value(row_idx, c) for c in range(sheet.ncols)]; study = _strip(row[camp_col])
                if not study or study.upper() == "TOTAL": continue
                swath = _strip(row[camp_col + 1]) if camp_col + 1 < len(row) else ""; n_raw = row[camp_col + 2] if camp_col + 2 < len(row) else ""
                data_fmt = _strip(row[camp_col + 3]) if camp_col + 3 < len(row) else ""; date_raw = row[camp_col + 4] if camp_col + 4 < len(row) else None
                nat_raw = _strip(row[camp_col + 5]) if camp_col + 5 < len(row) else ""; structure = _strip(row[0]) if has_structure else ""
                trailing = _join_obs(*[_strip(row[c]) for c in range(trailing_start, len(row)) if _strip(row[c])])
                obs = _xls_obs(structure, trailing, sheet_name); rec_date = _xldate(date_raw, book) if isinstance(date_raw, (int, float)) else None
                supp_type = ("Cassette" if nat_raw in ("3592.0", "3592", "3592,0") else _normalise_support(nat_raw or data_fmt))
                n_supp = max(_to_int(n_raw), 1) if n_raw != "" else 1
                for i in range(n_supp): records.append(make_record(study, rec_date, supp_type, i + 1, swath, data_fmt, "", obs, fname))
        return records

    def parse_sit1(self, filepath: Path) -> list:
        records = []; fname = filepath.name
        try: book = xlrd.open_workbook(str(filepath))
        except Exception as e: self.log_cb(f"⚠️ Cannot open {fname}: {e}"); return records
        for sheet_name in ("ASS", "SH"):
            if sheet_name not in book.sheet_names(): continue
            sheet = book.sheet_by_name(sheet_name); hdr1, hdr2 = 1, 2
            for r in range(min(10, sheet.nrows)):
                if any("structure" in _strip(sheet.cell_value(r, c)).lower() for c in range(sheet.ncols)): hdr1, hdr2 = r, r + 1; break
            sub_hdr = [_strip(sheet.cell_value(hdr2, c)).lower() for c in range(sheet.ncols)]
            cass_col = next((i for i, h in enumerate(sub_hdr) if "cass" in h), 6); cd_col = next((i for i, h in enumerate(sub_hdr) if "cd" in h), 7)
            doc_col = next((i for i, h in enumerate(sub_hdr) if "doc" in h), 8); nature_col = next((i for i, h in enumerate(sub_hdr) if "nature" in h or "supp" in h), 9)
            trailing_start = nature_col + 1
            for row_idx in range(hdr2 + 1, sheet.nrows):
                row = [sheet.cell_value(row_idx, c) for c in range(sheet.ncols)]; study = _strip(row[1]) if len(row) > 1 else ""; swath = _strip(row[2]) if len(row) > 2 else ""
                if not study or study.upper() in ("STRUCTURE", "ETUDE"): continue
                if "total" in _strip(row[3] if len(row) > 3 else "").lower(): continue
                n_cass = _to_int(row[cass_col]) if cass_col < len(row) else 0; n_cd = _to_int(row[cd_col]) if cd_col < len(row) else 0; n_doc = _to_int(row[doc_col]) if doc_col < len(row) else 0
                if n_cass == 0 and n_cd == 0 and n_doc == 0: continue
                structure = _strip(row[0]); trailing = _join_obs(*[_strip(row[c]) for c in range(trailing_start, len(row)) if _strip(row[c])])
                obs = _xls_obs(structure, trailing, sheet_name); records.extend(_emit_supports(n_cass, n_cd, 0, n_doc, study, None, swath, "", "", obs, fname))
            if records: break
        return records

    def parse_stack(self, filepath: Path) -> list:
        records = []; fname = filepath.name
        try: book = xlrd.open_workbook(str(filepath))
        except Exception as e: self.log_cb(f"⚠️ Cannot open {fname}: {e}"); return records
        for sheet_name in [s for s in ("SH", "ASS") if s in book.sheet_names()]:
            sheet = book.sheet_by_name(sheet_name); hdr1, hdr2 = 4, 5
            for r in range(min(10, sheet.nrows)):
                if any("structure" in _strip(sheet.cell_value(r, c)).lower() for c in range(sheet.ncols)): hdr1, hdr2 = r, r + 1; break
            main_hdr = [_strip(sheet.cell_value(hdr1, c)).lower() for c in range(sheet.ncols)]; sub_hdr = [_strip(sheet.cell_value(hdr2, c)).lower() for c in range(sheet.ncols)]
            cass_col = next((i for i, h in enumerate(sub_hdr) if "cass" in h), 4); cd_col = next((i for i, h in enumerate(sub_hdr) if "cd" in h), 5)
            usb_col = next((i for i, h in enumerate(sub_hdr) if "usb" in h), None); doc_col = next((i for i, h in enumerate(sub_hdr) if "doc" in h), None)
            centre_col = next((i for i, h in enumerate(main_hdr) if "centre" in h), None); date_col = next((i for i, h in enumerate(main_hdr) if "date" in h), 3)
            if sheet_name == "ASS":
                for c in range(sheet.ncols - 1, -1, -1):
                    if "date" in _strip(sheet.cell_value(hdr1, c)).lower(): date_col = c; break
            last_named = max([i for i, h in enumerate(main_hdr) if h] + [i for i, h in enumerate(sub_hdr) if h]); trailing_start = last_named + 1
            for row_idx in range(hdr2 + 1, sheet.nrows):
                row = [sheet.cell_value(row_idx, c) for c in range(sheet.ncols)]; study = _strip(row[1]) if len(row) > 1 else ""; swath = _strip(row[2]) if len(row) > 2 else ""
                if not study or study.upper() in ("STRUCTURE", "ETUDE"): continue
                if "total" in _strip(row[3] if len(row) > 3 else "").lower(): continue
                rec_date = _xldate(row[date_col], book) if date_col < len(row) and isinstance(row[date_col], (int, float)) else None
                centre = _strip(row[centre_col]) if centre_col and centre_col < len(row) else ""
                n_cass = _to_int(row[cass_col]) if cass_col < len(row) else 0; n_cd = _to_int(row[cd_col]) if cd_col < len(row) else 0
                n_usb = _to_int(row[usb_col]) if usb_col is not None and usb_col < len(row) else 0; n_doc = _to_int(row[doc_col]) if doc_col is not None and doc_col < len(row) else 0
                if n_cass == 0 and n_cd == 0 and n_usb == 0 and n_doc == 0: continue
                structure = _strip(row[0]); trailing = _join_obs(*[_strip(row[c]) for c in range(trailing_start, len(row)) if _strip(row[c])])
                obs = _xls_obs(structure, trailing, sheet_name); records.extend(_emit_supports(n_cass, n_cd, n_usb, n_doc, study, rec_date, swath, "stack", centre, obs, fname))
        return records

    def parse_reception(self, filepath: Path) -> list:
        records = []; fname = filepath.name
        try: wb = load_workbook(str(filepath), read_only=True, data_only=True)
        except Exception as e: self.log_cb(f"⚠️ Cannot open {fname}: {e}"); return records
        ws = wb.active; header = [ws.cell(1, c).value for c in range(1, (ws.max_column or 0) + 1)]
        kw_map = {"date": "date", "type": "data_type", "compagnie": "centre", "etude": "study", "profil": "swath", "cassette": "cass", "cd": "cd", "doc": "doc", "usb": "usb", "observ": "obs"}
        col_map = {}
        for idx, val in enumerate(header):
            h = _strip(val).lower()
            for kw, field in kw_map.items():
                if kw in h and field not in col_map: col_map[field] = idx; break
        missing = {"study", "cass", "cd"} - col_map.keys()
        if missing: self.log_cb(f"⚠️ {fname}: Missing {missing}"); wb.close(); return records
        def _get(row_vals, key, default=None): idx = col_map.get(key); return row_vals[idx] if (idx is not None and idx < len(row_vals)) else default
        for row in ws.iter_rows(min_row=2, values_only=True):
            row = list(row); study = _strip(_get(row, "study"))
            if not study: continue
            rec_date = _to_date(_get(row, "date")); swath = _strip(_get(row, "swath")); data_type = _strip(_get(row, "data_type"))
            centre = _strip(_get(row, "centre")); obs = _strip(_get(row, "obs"))
            n_cass = _to_int(_get(row, "cass", 0)); n_cd = _to_int(_get(row, "cd", 0)); n_usb = _to_int(_get(row, "usb", 0)); n_doc = _to_int(_get(row, "doc", 0))
            if n_cass == 0 and n_cd == 0 and n_usb == 0 and n_doc == 0: continue
            records.extend(_emit_supports(n_cass, n_cd, n_usb, n_doc, study, rec_date, swath, data_type, centre, obs, fname))
        wb.close(); return records

    def parse_template(self, filepath: Path) -> list:
        records = []; fname = filepath.name
        try: book = xlrd.open_workbook(str(filepath))
        except Exception as e: self.log_cb(f"⚠️ Cannot open {fname}: {e}"); return records
        sheet = book.sheets()[0]
        section_rows = [r for r in range(sheet.nrows) if "directions" in " ".join(_strip(sheet.cell_value(r, c)).lower() for c in range(sheet.ncols))]
        if not section_rows: return records
        for idx, hdr_row in enumerate(section_rows):
            next_hdr = section_rows[idx + 1] if idx + 1 < len(section_rows) else sheet.nrows
            row_text = " ".join(_strip(sheet.cell_value(hdr_row, c)).lower() for c in range(sheet.ncols))
            if "date de rcep" in row_text:
                for row_idx in range(hdr_row + 2, next_hdr):
                    row = [sheet.cell_value(row_idx, c) for c in range(sheet.ncols)]; study = _strip(row[1]) if len(row) > 1 else ""
                    if not study or any("total" in v.lower() for v in [_strip(v) for v in row if _strip(v)]): continue
                    n_cass = _to_int(row[6]) if len(row) > 6 else 0; n_cd = _to_int(row[7]) if len(row) > 7 else 0; n_doc = _to_int(row[8]) if len(row) > 8 else 0
                    if n_cass == 0 and n_cd == 0 and n_doc == 0: continue
                    records.extend(_emit_supports(n_cass, n_cd, 0, n_doc, study, _to_date(row[9] if len(row) > 9 else None, book), _strip(row[2]) if len(row) > 2 else "", "terrain", "", _join_obs(f"Structure : {_strip(row[0])}" if len(row) > 0 and _strip(row[0]) and len(_strip(row[0])) > 1 else "", _strip(row[11]) if len(row) > 11 else ""), fname))
            elif any(k in row_text for k in ("date de r", "réception", "reception")):
                main_hdr = [_strip(sheet.cell_value(hdr_row, c)).lower() for c in range(sheet.ncols)]; sub_hdr = [_strip(sheet.cell_value(hdr_row + 1, c)).lower() for c in range(sheet.ncols)]
                date_col = next((i for i, h in enumerate(main_hdr) if "date" in h), 3); cass_col = next((i for i, h in enumerate(sub_hdr) if "cass" in h), 4)
                cd_col = next((i for i, h in enumerate(sub_hdr) if "cd" in h), 5); doc_col = next((i for i, h in enumerate(sub_hdr) if "doc" in h), 6)
                usb_col = next((i for i, h in enumerate(sub_hdr) if "usb" in h), 7); centre_col = next((i for i, h in enumerate(main_hdr) if "centre" in h), 9)
                for row_idx in range(hdr_row + 2, next_hdr):
                    row = [sheet.cell_value(row_idx, c) for c in range(sheet.ncols)]; study = _strip(row[1]) if len(row) > 1 else ""
                    if not study or any("total" in v.lower() for v in [_strip(v) for v in row if _strip(v)]): continue
                    n_cass = _to_int(row[cass_col]) if cass_col < len(row) else 0; n_cd = _to_int(row[cd_col]) if cd_col < len(row) else 0
                    n_doc = _to_int(row[doc_col]) if doc_col < len(row) else 0; n_usb = _to_int(row[usb_col]) if usb_col < len(row) else 0
                    if n_cass == 0 and n_cd == 0 and n_doc == 0 and n_usb == 0: continue
                    records.extend(_emit_supports(n_cass, n_cd, n_usb, n_doc, study, _to_date(row[date_col] if date_col < len(row) else None, book), _strip(row[2]) if len(row) > 2 else "", "stack", _strip(row[centre_col]) if centre_col < len(row) else "", _join_obs(f"Structure : {_strip(row[0])}" if len(row) > 0 and _strip(row[0]) and len(_strip(row[0])) > 1 else ""), fname))
        return records

    def parse_file(self, filepath: Path) -> list:
        family = self.detect_family(filepath)
        if family == FAMILY_SIT2: return self.parse_sit2(filepath)
        if family == FAMILY_SIT1: return self.parse_sit1(filepath)
        if family == FAMILY_STACK: return self.parse_stack(filepath)
        if family == FAMILY_RECEPTION: return self.parse_reception(filepath)
        if family == FAMILY_TEMPLATE: return self.parse_template(filepath)
        return []

    def _capture_row_style(self, ws, row_num: int) -> dict:
        styles = {}
        for cell in next(ws.iter_rows(min_row=row_num, max_row=row_num)):
            styles[cell.column] = {"font": copy.copy(cell.font) if cell.font else None, "border": copy.copy(cell.border) if cell.border else None, "alignment": copy.copy(cell.alignment) if cell.alignment else None, "number_format": cell.number_format}
        return styles

    def write_output(self, records: list, template_path: Path, output_path: Path) -> None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        wb = load_workbook(str(template_path)); ws = wb.active
        try: row_style = self._capture_row_style(ws, 3)
        except: row_style = {}
        write_row = 2
        for row_idx in range(ws.max_row, 1, -1):
            if ws.cell(row=row_idx, column=1).value is not None: write_row = row_idx + 1; break
        col_map = {name: idx + 1 for idx, name in enumerate(OUTPUT_COLUMNS)}
        k_col = col_map["source_file"]
        if _strip(ws.cell(1, k_col).value).lower() != "source_file":
            hdr_cell = ws.cell(1, k_col); hdr_cell.value = "source_file"; ref = ws.cell(1, 1)
            hdr_cell.font = copy.copy(ref.font); hdr_cell.fill = copy.copy(ref.fill); hdr_cell.border = copy.copy(ref.border); hdr_cell.alignment = copy.copy(ref.alignment)
        self.log_cb(f"> Writing {len(records)} rows starting at row {write_row}...")
        for rec in records:
            fill = SUPPORT_COLOURS.get(_strip(rec.support_type), SUPPORT_COLOURS["Unknown"])
            for col_name, col_idx in col_map.items():
                cell = ws.cell(row=write_row, column=col_idx); cell.value = getattr(rec, col_name); cell.fill = fill
                if col_idx in row_style:
                    st = row_style[col_idx]
                    if st["font"]: cell.font = copy.copy(st["font"])
                    if st["border"]: cell.border = copy.copy(st["border"])
                    if st["alignment"]: cell.alignment = copy.copy(st["alignment"])
                    if st["number_format"]: cell.number_format = st["number_format"]
                if col_name == "reception_date" and isinstance(cell.value, date): cell.number_format = "DD/MM/YYYY"
            write_row += 1
        wb.save(str(output_path))
        self.log_cb(f"> Saved master workbook to {output_path.name}")

# ==========================================
# FLASK ROUTES
# ==========================================
@app.route('/')
def index(): return render_template('index.html')

@app.route('/synthese')
def synthese(): return render_template('synthese.html')

@app.route('/api/drive/list', methods=['POST'])
def api_list():
    data = request.get_json(); lines = list_contents(data.get('path'), data.get('show_hidden', False), data.get('show_sizes', True), data.get('recursive', False))
    if data.get('export'):
        mem = io.BytesIO("\n".join(lines).encode('utf-8')); mem.seek(0)
        return send_file(mem, as_attachment=True, download_name=f"listing_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt", mimetype='text/plain')
    return jsonify({"lines": lines})

@app.route('/api/drive/compare', methods=['POST'])
def api_compare():
    data = request.get_json(); lines = compare_listings(data.get('path_a'), data.get('path_b'), data.get('show_hidden', False))
    if data.get('export'):
        mem = io.BytesIO("\n".join(lines).encode('utf-8')); mem.seek(0)
        return send_file(mem, as_attachment=True, download_name=f"compare_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt", mimetype='text/plain')
    return jsonify({"lines": lines})

@app.route('/api/drive/compare-files', methods=['POST'])
def api_compare_files():
    f_a, f_b = request.files.get('file_a'), request.files.get('file_b')
    if not f_a or not f_b: return jsonify({"error": "Both files required"}), 400
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as t_a, tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as t_b:
        f_a.save(t_a.name); f_b.save(t_b.name)
        p_a, p_b = parse_listing_file(t_a.name), parse_listing_file(t_b.name)
        os.unlink(t_a.name); os.unlink(t_b.name)
    lines = compare_parsed_listings(p_a, p_b)
    if request.form.get('export') == 'true':
        mem = io.BytesIO("\n".join(lines).encode('utf-8')); mem.seek(0)
        return send_file(mem, as_attachment=True, download_name=f"txt_compare_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt", mimetype='text/plain')
    return jsonify({"lines": lines})

@app.route('/api/drive/roots', methods=['GET'])
def get_roots():
    roots = []
    if platform.system() == 'Windows':
        for letter in string.ascii_uppercase:
            if os.path.exists(f"{letter}:\\"): roots.append(f"{letter}:\\")
    else:
        roots.append('/')
        for p in ['/media', '/mnt']:
            if os.path.exists(p): roots.extend([f"{p}/{d}" for d in os.listdir(p)])
    return jsonify({"roots": roots})

@app.route('/api/drive/folders', methods=['POST'])
def get_folders():
    data = request.get_json(); target_path = data.get('path', '').strip()
    if not target_path or not os.path.exists(target_path) or not os.path.isdir(target_path): return jsonify({"error": "Invalid path"}), 400
    try:
        dirs = [{"name": entry.name, "full_path": entry.path} for entry in os.scandir(target_path) if entry.is_dir() and entry.name not in ['$RECYCLE.BIN', 'System Volume Information']]
        dirs.sort(key=lambda x: x['name'].lower()); parent = os.path.dirname(target_path)
        return jsonify({"current": target_path, "parent": parent if parent != target_path else None, "dirs": dirs})
    except Exception as e: return jsonify({"error": str(e)}), 500

@app.route('/api/pdf/compress', methods=['POST'])
def api_compress_pdf():
    if pdfplumber is None: return jsonify({"error": "pdfplumber not installed"}), 500
    file = request.files['file']; aggressive = request.form.get('aggressive') == 'true'
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
        file.save(temp_pdf.name); temp_path = temp_pdf.name
    try:
        original_size = os.path.getsize(temp_path); all_pages = []
        with pdfplumber.open(temp_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_md = extract_page_content(page, aggressive)
                if page_md.strip(): all_pages.append(f"\n\n<!-- Page {i} -->\n\n{page_md}")
        full_md = re.sub(r"\n{3,}", "\n\n", WHITESPACE_PATTERN.sub(" ", "\n".join(all_pages).strip()))
        full_md = f"# {os.path.splitext(file.filename)[0]}\n> **Mode:** {'Aggressive' if aggressive else 'Standard'}\n---\n" + full_md
        compressed_size = len(full_md.encode('utf-8'))
        return jsonify({"markdown": full_md, "original_kb": round(original_size / 1024, 2), "compressed_kb": round(compressed_size / 1024, 2), "reduction_percent": round(((original_size - compressed_size) / original_size) * 100, 1)})
    except Exception as e: return jsonify({"error": str(e)}), 500
    finally: os.unlink(temp_path)

# ==========================================
# APP 2: SEISMIC MINER ROUTES
# ==========================================
@app.route('/api/miner/stream', methods=['POST'])
def miner_stream():
    template_file = request.files.get('template')
    source_files = request.files.getlist('sources')
    
    temp_dir = Path(tempfile.mkdtemp())
    output_id = str(uuid.uuid4())
    output_filename = f"receptions_filled_{output_id}.xlsx"
    output_dir = Path("outputs")
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / output_filename

    template_filename = os.path.basename(template_file.filename)
    template_path = temp_dir / template_filename
    template_file.save(str(template_path))
    
    for f in source_files:
        source_filename = os.path.basename(f.filename)
        f.save(str(temp_dir / source_filename))

    log_queue = queue.Queue()

    def run_task():
        try:
            log_queue.put(f"> Template loaded: {template_path.name}")
            log_queue.put(f"> {len(source_files)} source files staged in temp environment.")

            consolidator = SeismicConsolidator(log_queue.put)
            
            try:
                rec_template = consolidator.find_reception_template(temp_dir, output_path)
                log_queue.put(f"> Output template identified: {rec_template.name}")
            except FileNotFoundError as e:
                log_queue.put(f"❌ Error: {str(e)}")
                log_queue.put("DONE_ERROR"); return

            sources = consolidator.discover_sources(temp_dir, output_path)
            log_queue.put(f"> Discovered {len(sources)} processable source files.")
            if not sources: 
                log_queue.put("❌ Error: No valid source files found.")
                log_queue.put("DONE_ERROR"); return

            all_records = []
            for filepath in sources:
                family = consolidator.detect_family(filepath)
                log_queue.put(f"> Parsing {filepath.name} ({family})...")
                try: 
                    all_records.extend(consolidator.parse_file(filepath))
                except Exception as exc: 
                    log_queue.put(f"⚠️ Error parsing {filepath.name}: {exc}")

            log_queue.put(f"> Total records extracted: {len(all_records)}")
            if not all_records: 
                log_queue.put("❌ Error: No records extracted. Aborting.")
                log_queue.put("DONE_ERROR"); return

            consolidator.write_output(all_records, rec_template, output_path)
            log_queue.put(f"✅ Pipeline completed successfully.")
            log_queue.put(f"DOWNLOAD_URL:/api/miner/download/{output_id}")
            log_queue.put("DONE_SUCCESS")
        except Exception as e:
            log_queue.put(f"❌ FATAL: {str(e)}")
            log_queue.put("DONE_ERROR")
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    threading.Thread(target=run_task).start()

    def generate():
        while True:
            try:
                msg = log_queue.get(timeout=30)
                if msg == "DONE_SUCCESS": 
                    yield f"data: {json.dumps({'log': '✅ Done', 'done': True, 'success': True})}\n\n"
                    break
                elif msg == "DONE_ERROR": 
                    yield f"data: {json.dumps({'log': '❌ Failed', 'done': True, 'success': False})}\n\n"
                    break
                elif str(msg).startswith("DOWNLOAD_URL:"): 
                    yield f"data: {json.dumps({'log': 'Generating download link...', 'download_url': str(msg).split(':', 1)[1]})}\n\n"
                else: 
                    yield f"data: {json.dumps({'log': str(msg)})}\n\n"
            except queue.Empty: 
                yield f"data: {json.dumps({'log': '...'})}\n\n"

    return Response(generate(), mimetype='text/event-stream')

@app.route('/api/miner/download/<file_id>')
def miner_download(file_id):
    filepath = Path("outputs") / f"receptions_filled_{file_id}.xlsx"
    if filepath.exists(): return send_file(filepath, as_attachment=True, download_name="receptions_filled.xlsx")
    return "File not found", 404



# ==========================================
# MONTHLY REPORTS FUNCTIONS & ROUTES
# ==========================================
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Sonatrach Situation Package E&P — Report Generator Web App
Flask backend with full report generation logic
"""


import copy
import io
import logging
import logging.handlers
import os
import re
import unicodedata
import uuid
from dataclasses import dataclass
from datetime import datetime, timedelta
from difflib import SequenceMatcher
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from flask import Flask, jsonify, render_template, request, send_file



UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")
LOG_DIR    = Path("logs")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)
LOG_DIR.mkdir(exist_ok=True)

# ─────────────────────────────────────────────────────────────────────────────
# LOGGING — rotating file log of every upload/validate/generate/download, plus
# console output. No file contents are logged, only filenames/ids/summaries,
# so this is safe to keep around for troubleshooting without becoming a data
# store of its own.
# ─────────────────────────────────────────────────────────────────────────────

logger = logging.getLogger("sonatrach_report")
logger.setLevel(logging.INFO)
if not logger.handlers:
    _file_handler = logging.handlers.RotatingFileHandler(
        LOG_DIR / "app.log", maxBytes=2_000_000, backupCount=5, encoding="utf-8"
    )
    _console_handler = logging.StreamHandler()
    _fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
    _file_handler.setFormatter(_fmt)
    _console_handler.setFormatter(_fmt)
    logger.addHandler(_file_handler)
    logger.addHandler(_console_handler)


# ─────────────────────────────────────────────────────────────────────────────
# ERROR TYPES — used to return specific, actionable messages to the UI
# instead of a generic "Internal Server Error" for anything that goes wrong.
# ─────────────────────────────────────────────────────────────────────────────


class AppError(Exception):
    """Base class for errors we can explain clearly to the user."""
    status_code = 400

    def __init__(self, message: str, status_code: int | None = None):
        super().__init__(message)
        self.message = message
        if status_code is not None:
            self.status_code = status_code


class FileNotFoundInStoreError(AppError):
    status_code = 404


class ExcelFormatError(AppError):
    status_code = 422


class TemplateFormatError(AppError):
    status_code = 422

# ─────────────────────────────────────────────────────────────────────────────
# HEADING RULES
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class HeadingRule:
    category: str
    label: str
    must_contain: tuple
    is_main_section: bool = False


HEADING_RULES: list[HeadingRule] = [
    HeadingRule("PED",        "VI-1- Package de données E&P",           ("vi-1-", "package", "donnees"), is_main_section=True),
    HeadingRule("EDS",        "VI-1-2 Package_EDS",                     ("vi-1-2", "eds")),
    HeadingRule("Asset",      "VI-1-3 Package_Asset",                   ("vi-1-3", "asset")),
    HeadingRule("BDD",        "VI-1-4 Package BDD",                     ("vi-1-4", "bdd")),
    HeadingRule("Traitement", "VI-1-5 Package Traitement/Retraitement", ("vi-1-5",)),
    HeadingRule("SINOPEC",    "VI-1-6 Package_SINOPEC",                 ("vi-1-6", "sinopec")),
    HeadingRule("ENI",        "VI-1-7 Package ENI",                     ("vi-1-7", "eni")),
    HeadingRule("ZPEC",       "VI-1-8 Package ZPEC",                    ("vi-1-8", "zpec")),
    HeadingRule("Total",      "VI-1-9 Package Total",                   ("vi-1-9",)),
    HeadingRule("SH & ZANGAZ&FILADA", "VI-1-10 SH & ZANGAZ&FILADA",     ("vi-1-10",)),
]

CATEGORY_LABELS = {r.category: r.label for r in HEADING_RULES}

# For most categories, the Excel "BDD" column value is just the category
# name itself (e.g. "PED" rows -> PED table). "Asset" is the one exception:
# its table is filled from FOUR separate BDD values, not a single "Asset"
# row. Every other category keeps matching on its own name exactly as before.
CATEGORY_BDD_VALUES: dict[str, list[str]] = {
    "Asset": ["ASCentre", "ASOuest", "ASEst", "ASNord"],
    "Traitement": ["DP/DMG"],
}


def bdd_values_for_category(category: str) -> list[str]:
    return CATEGORY_BDD_VALUES.get(category, [category])

COLUMN_KEYWORDS: dict[str, tuple] = {
    "etudes":      ("etude",),
    "type":        ("type",),
    "date":        ("date",),
    "source":      ("source",),
    "capacite":    ("capacite",),
    "realisation": ("realisation",),
    "remarque":    ("remarque",),
}

EXPECTED_COLUMNS = ["BDD", "Data", "Type_de_donnees", "Date", "Source", "Capacite", "Realisation", "Remarque"]

FIELD_TO_COLUMN = {
    "etudes":      "Data_fmt",
    "type":        "Type_fmt",
    "date":        "Date_fmt",
    "source":      "Source_fmt",
    "capacite":    "Capacite_fmt",
    "realisation": "Realisation_fmt",
    "remarque":    "Remarque_fmt",
}

# ─────────────────────────────────────────────────────────────────────────────
# NORMALISATION HELPERS
# ─────────────────────────────────────────────────────────────────────────────

_DASH_VARIANTS = "\u2010\u2011\u2012\u2013\u2014\u2015\u2212"  # various unicode dash characters


def normalize(text: str) -> str:
    if text is None:
        return ""
    text = unicodedata.normalize("NFKD", str(text))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower()
    for dash in _DASH_VARIANTS:
        text = text.replace(dash, "-")
    text = re.sub(r"[_\u00a0]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def heading_matches(norm_text: str, rule: HeadingRule) -> bool:
    if not norm_text:
        return False
    if rule.is_main_section:
        if not re.match(r"^vi-1-\s", norm_text) and not re.match(r"^vi-1-[^0-9]", norm_text):
            return False
        if re.match(r"^vi-1-\d", norm_text):
            return False
    return all(token in norm_text for token in rule.must_contain)


def study_key(text) -> str:
    """A robust matching key for Étude / Permis codes (e.g. '2025-DJA-2D').
    Strips ALL separators, parenthetical annotations, and casing so that
    formatting differences that don't change the actual code — hyphen vs
    space vs underscore, extra whitespace, a trailing '(Zeml)' style note —
    never cause a genuine match to be missed. Used only for Étude/Permis
    matching, never for heading detection (which needs '-' preserved)."""
    if text is None:
        return ""
    t = normalize(text)
    t = re.sub(r"\([^)]*\)", " ", t)   # drop parenthetical annotations
    t = re.sub(r"[^a-z0-9]", "", t)    # keep only letters/digits
    return t


def study_key_sorted(text) -> str:
    """Order-invariant sibling of study_key(): same normalization, but the
    individual alphanumeric tokens are sorted before joining. This exists
    for one specific, safe purpose — a trailing descriptor written in a
    different word order, e.g. '2025-HMD-Sud-3D' vs '2025-HMD-3D SUD',
    should still count as the same étude. Since it compares the exact same
    multiset of tokens (not a character-similarity score), it can't drift
    into the kind of coincidental match study_key_similarity() warns about:
    two genuinely different codes never share every token, only their
    order. study_key() itself is left untouched — containment matching
    (tier 2) still needs the original left-to-right order preserved so an
    annotation appended after the real code (e.g. '(Zeml)') keeps working."""
    if text is None:
        return ""
    t = normalize(text)
    t = re.sub(r"\([^)]*\)", " ", t)
    tokens = re.findall(r"[a-z0-9]+", t)
    return "".join(sorted(tokens))


def study_key_similarity(a: str, b: str) -> float:
    """0..1 similarity between two study_key() outputs. IMPORTANT: this is
    used only to surface a "closest candidate" hint for a human to review —
    it is NEVER used to automatically decide a match. Étude/Permis codes in
    this dataset (e.g. '2025-DJA-2D', '2026-EAGLE-2D', '2025-YOUCEF-3D') all
    share the same 'YYYY-XXX-ND' shape, so two genuinely different codes
    routinely score 40-60% similar on generic character overlap alone —
    auto-applying at that kind of threshold silently writes real reception
    data (profiles, cassette counts) into the wrong étude's row. See
    MIN_CONTAINMENT_LEN below for the (safe) leniency this script actually
    applies automatically."""
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()


# Minimum length of a study_key() for it to be eligible for the "containment"
# match tier (tier 2) below — e.g. matching '2025ZAR3D' inside
# '2025zar3dzeml'. Below this length, a substring match is too likely to be
# coincidental (e.g. a bare "3d" or "2d") to trust automatically.
MIN_CONTAINMENT_LEN = 6


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT WALKING
# ─────────────────────────────────────────────────────────────────────────────

def iter_block_items(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def locate_tables(document: Document) -> dict[str, Table]:
    assigned: dict[str, Table] = {}
    pending: Optional[HeadingRule] = None

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            norm = normalize(block.text)
            if not norm:
                continue
            for rule in HEADING_RULES:
                if rule.category in assigned:
                    continue
                if heading_matches(norm, rule):
                    pending = rule
                    break
        elif isinstance(block, Table):
            if pending is not None and pending.category not in assigned:
                assigned[pending.category] = block
                pending = None

    return assigned


def locate_table_columns(table: Table) -> dict[str, int]:
    header_cells = table.rows[0].cells
    mapping: dict[str, int] = {}
    for idx, cell in enumerate(header_cells):
        norm = normalize(cell.text)
        for field_name, keywords in COLUMN_KEYWORDS.items():
            if field_name in mapping:
                continue
            if any(kw in norm for kw in keywords):
                mapping[field_name] = idx
    return mapping


# ─────────────────────────────────────────────────────────────────────────────
# TEMPLATE VALIDATION — pre-flight check run right after a Word template is
# uploaded, so a broken/edited template surfaces a clear diagnostic instead
# of a silent "table not found" (or worse, a wrong table) at generation time.
# ─────────────────────────────────────────────────────────────────────────────

def _vi1_heading_found_map(document: Document) -> dict[str, bool]:
    found = {r.category: False for r in HEADING_RULES}
    for block in iter_block_items(document):
        if not isinstance(block, Paragraph):
            continue
        norm = normalize(block.text)
        if not norm:
            continue
        for rule in HEADING_RULES:
            if not found[rule.category] and heading_matches(norm, rule):
                found[rule.category] = True
    return found


def _vi3_heading_found(document: Document) -> tuple[bool, bool]:
    terrain_found = False
    stack_found = False
    for block in iter_block_items(document):
        if not isinstance(block, Paragraph):
            continue
        norm = normalize(block.text)
        if not norm:
            continue
        if "reception" in norm and "sismique" in norm and "terrain" in norm:
            terrain_found = True
        elif "reception" in norm and "sismique" in norm and "stack" in norm:
            stack_found = True
    return terrain_found, stack_found


def _table_missing_columns(table: Table, expected: dict[str, tuple]) -> list[str]:
    missing = []
    for field, tokens in expected.items():
        if _find_table_col(table, tokens) is None:
            missing.append(field)
    return missing


VI3_TERRAIN_EXPECTED_COLUMNS = {
    "permis":          ("permis",),
    "donnees_recues":  ("donnees", "recu"),
    "nombre_sup_3592": ("sup",),
}

VI3_STACK_EXPECTED_COLUMNS = {
    "direction":      ("direction",),
    "projet":         ("projet",),
    "centre":         ("centre",),
    "date_reception": ("date",),
    "donnees_recues": ("donnees", "recu"),
}


def validate_template(document: Document) -> dict:
    """Run a full structural check of the Word template: every VI-1 heading,
    every VI-3 heading, whether a table follows each, and whether the
    expected columns are present in each table found. Never raises — always
    returns a report, so it's safe to call speculatively right after upload."""

    vi1_heading_found = _vi1_heading_found_map(document)
    vi1_table_map = locate_tables(document)

    vi1_report: dict[str, dict] = {}
    for rule in HEADING_RULES:
        cat = rule.category
        table = vi1_table_map.get(cat)
        entry = {
            "label": rule.label,
            "heading_found": vi1_heading_found[cat],
            "table_found": table is not None,
            "missing_columns": None,
            "data_rows": None,
        }
        if table is not None:
            col_map = locate_table_columns(table)
            entry["missing_columns"] = [f for f in COLUMN_KEYWORDS if f not in col_map]
            entry["data_rows"] = len(table.rows) - 1

        if not entry["heading_found"]:
            entry["status"] = "missing"
        elif not entry["table_found"]:
            entry["status"] = "error"
        elif entry["missing_columns"]:
            entry["status"] = "warning"
        else:
            entry["status"] = "ok"
        vi1_report[cat] = entry

    terrain_heading_found, stack_heading_found = _vi3_heading_found(document)
    reception_tables = locate_reception_tables(document)

    terrain_tables = reception_tables.get("terrain", [])
    terrain_table_reports = [
        {"missing_columns": _table_missing_columns(t, VI3_TERRAIN_EXPECTED_COLUMNS), "data_rows": len(t.rows) - 1}
        for t in terrain_tables
    ]
    vi3_terrain = {
        "heading_found": terrain_heading_found,
        "tables_found":  len(terrain_tables),
        "tables":        terrain_table_reports,
    }
    if not terrain_heading_found:
        vi3_terrain["status"] = "missing"
    elif not terrain_tables:
        vi3_terrain["status"] = "error"
    elif any(t["missing_columns"] for t in terrain_table_reports):
        vi3_terrain["status"] = "warning"
    else:
        vi3_terrain["status"] = "ok"

    stack_tables = reception_tables.get("stack", [])
    stack_table_reports = [
        {"missing_columns": _table_missing_columns(t, VI3_STACK_EXPECTED_COLUMNS), "data_rows": len(t.rows) - 1}
        for t in stack_tables
    ]
    vi3_stack = {
        "heading_found": stack_heading_found,
        "tables_found":  len(stack_tables),
        "tables":        stack_table_reports,
    }
    if not stack_heading_found:
        vi3_stack["status"] = "missing"
    elif not stack_tables:
        vi3_stack["status"] = "error"
    elif any(t["missing_columns"] for t in stack_table_reports):
        vi3_stack["status"] = "warning"
    else:
        vi3_stack["status"] = "ok"

    all_statuses = [e["status"] for e in vi1_report.values()] + [vi3_terrain["status"], vi3_stack["status"]]
    if "error" in all_statuses:
        overall = "error"
    elif "warning" in all_statuses:
        overall = "warning"
    else:
        overall = "ok"

    return {
        "vi1":            vi1_report,
        "vi3_terrain":    vi3_terrain,
        "vi3_stack":      vi3_stack,
        "overall_status": overall,
    }


# ─────────────────────────────────────────────────────────────────────────────
# EXCEL READING
# ─────────────────────────────────────────────────────────────────────────────

def excel_date(val) -> str:
    if pd.isna(val) or val == "":
        return ""
    if isinstance(val, datetime):
        return val.strftime("%d/%m/%Y")
    if isinstance(val, (int, float)):
        try:
            return (datetime(1899, 12, 30) + timedelta(days=int(val))).strftime("%d/%m/%Y")
        except Exception:
            return str(val)
    if isinstance(val, str):
        for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%m/%d/%Y"):
            try:
                return datetime.strptime(val.strip(), fmt).strftime("%d/%m/%Y")
            except ValueError:
                pass
        return val.strip()
    return str(val)


def excel_number(val) -> str:
    if pd.isna(val) or val == "":
        return ""
    try:
        f = float(val)
        return str(int(f)) if f == int(f) else str(round(f, 4))
    except (ValueError, TypeError):
        return str(val).strip()


def excel_multiline(val) -> str:
    if pd.isna(val) or val == "":
        return ""
    s = str(val).strip()
    s = re.sub(r"\s*,\s*", "\n", s)
    return s


def excel_text(val) -> str:
    if pd.isna(val) or val == "":
        return ""
    return str(val).strip()


def read_excel_data(excel_path: Path) -> pd.DataFrame:
    try:
        raw = pd.read_excel(excel_path, header=None)
    except Exception as e:
        raise ExcelFormatError(
            f"Could not open '{excel_path.name}' as an Excel file. "
            f"Make sure it's a valid .xlsx/.xls file and isn't corrupted or password-protected. ({e})"
        )

    if raw.empty:
        raise ExcelFormatError(f"'{excel_path.name}' appears to be empty.")

    header_row_idx = None
    for i in range(min(len(raw), 25)):
        first_cell = str(raw.iat[i, 0]).strip().lower() if pd.notna(raw.iat[i, 0]) else ""
        if first_cell == "bdd":
            header_row_idx = i
            break

    if header_row_idx is None:
        raise ExcelFormatError(
            "Could not find the header row (a row starting with 'BDD') in the Excel file. "
            "Check that this is the 'Situation Package' export and that its layout hasn't changed."
        )

    df = raw.iloc[header_row_idx + 1:].reset_index(drop=True)
    n_cols = len(EXPECTED_COLUMNS)
    if df.shape[1] < n_cols:
        raise ExcelFormatError(
            f"Expected at least {n_cols} columns starting from the 'BDD' header row, "
            f"but only found {df.shape[1]}. The Excel layout may have changed."
        )
    df = df.iloc[:, :n_cols]
    df.columns = EXPECTED_COLUMNS

    df = df.dropna(subset=["BDD"])
    df["BDD"] = df["BDD"].astype(str).str.strip()
    df = df[df["BDD"] != "nan"]

    if df.empty:
        raise ExcelFormatError(
            "No usable data rows were found under the 'BDD' header row. "
            "Check that the Excel file has been filled in."
        )

    df["Data_fmt"]        = df["Data"].apply(excel_multiline)
    df["Type_fmt"]        = df["Type_de_donnees"].apply(excel_text)
    df["Date_fmt"]        = df["Date"].apply(excel_date)
    df["Source_fmt"]      = df["Source"].apply(excel_text)
    df["Capacite_fmt"]    = df["Capacite"].apply(excel_number)
    df["Realisation_fmt"] = df["Realisation"].apply(excel_date)
    df["Remarque_fmt"]    = df["Remarque"].apply(excel_text)

    # Store raw date for month filtering
    df["_raw_date"] = pd.to_datetime(df["Date"], errors="coerce")

    return df


def get_available_months(df: pd.DataFrame) -> list[dict]:
    """Return list of {year, month, label} dicts sorted chronologically."""
    months_seen = set()
    for _, row in df.iterrows():
        d = row["_raw_date"]
        if pd.notna(d):
            months_seen.add((d.year, d.month))
    result = []
    for y, m in sorted(months_seen):
        label = datetime(y, m, 1).strftime("%B %Y")
        result.append({"year": y, "month": m, "label": label})
    return result


def filter_by_month(df: pd.DataFrame, year: int, month: int) -> pd.DataFrame:
    mask = df["_raw_date"].apply(
        lambda d: pd.notna(d) and d.year == year and d.month == month
    )
    return df[mask].copy()

# ─────────────────────────────────────────────────────────────────────────────
# CELL / ROW WRITING (format-preserving)
# ─────────────────────────────────────────────────────────────────────────────

def _clear_run_content(run):
    r = run._r
    for child in list(r):
        if child.tag in (qn("w:t"), qn("w:br"), qn("w:tab")):
            r.remove(child)


def _write_lines_into_run(run, lines: list[str]):
    _clear_run_content(run)
    r = run._r
    for i, line in enumerate(lines):
        if i > 0:
            r.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r.append(t)


def _ensure_run_with_format(paragraph, reference_run=None):
    if paragraph.runs:
        for extra in paragraph.runs[1:]:
            extra._r.getparent().remove(extra._r)
        return paragraph.runs[0]
    new_run = paragraph.add_run()
    if reference_run is not None:
        ref_rpr = reference_run._r.find(qn("w:rPr"))
        if ref_rpr is not None:
            new_run._r.insert(0, copy.deepcopy(ref_rpr))
    return new_run


def set_cell_text(cell: _Cell, text: str, reference_cell: Optional[_Cell] = None):
    paragraphs = cell.paragraphs
    for extra_para in paragraphs[1:]:
        extra_para._p.getparent().remove(extra_para._p)
    target_para = cell.paragraphs[0]

    reference_run = None
    if reference_cell is not None and reference_cell.paragraphs and reference_cell.paragraphs[0].runs:
        reference_run = reference_cell.paragraphs[0].runs[0]

    run = _ensure_run_with_format(target_para, reference_run)
    lines = text.split("\n") if text else [""]
    _write_lines_into_run(run, lines)


def clone_row(table: Table, template_row_index: int):
    template_tr = table.rows[template_row_index]._tr
    new_tr = copy.deepcopy(template_tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def remove_row(table: Table, row_index: int):
    tr = table.rows[row_index]._tr
    table._tbl.remove(tr)


def fill_table(table: Table, rows_data: list[dict], category: str) -> int:
    col_map = locate_table_columns(table)
    if not col_map:
        return 0

    header_row = table.rows[0]
    existing_data_rows = list(table.rows[1:])
    needed = len(rows_data)

    # Keep one data row as a formatting template before any trimming/clearing
    # Use the first existing data row (not the header) so we inherit correct
    # font/color/size (e.g. Calibri black) rather than the header style (bold white Arial).
    format_template_row_idx = 1 if len(table.rows) > 1 else 0

    if needed > len(existing_data_rows):
        while len(existing_data_rows) < needed:
            new_row = clone_row(table, format_template_row_idx)
            existing_data_rows.append(new_row)
    elif needed < len(existing_data_rows):
        surplus = len(existing_data_rows) - needed
        for _ in range(surplus):
            remove_row(table, len(table.rows) - 1)
        existing_data_rows = existing_data_rows[:needed]

    # Use the first data row as the formatting reference for new/empty cells.
    # This ensures cloned rows inherit Calibri/Arial black (data style), not
    # the bold white Arial of the header row.
    ref_row = existing_data_rows[0] if existing_data_rows else header_row

    written = 0
    for row_obj, record in zip(existing_data_rows, rows_data):
        for field_name, col_idx in col_map.items():
            if col_idx < len(row_obj.cells):
                cell = row_obj.cells[col_idx]
                ref_cell = ref_row.cells[col_idx]
                value = record.get(field_name, "")
                set_cell_text(cell, value, reference_cell=ref_cell)
        written += 1

    return written

# ─────────────────────────────────────────────────────────────────────────────
# MAIN GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_report(excel_path: Path, word_path: Path, year: int, month: int,
                    selected_categories: list[str],
                    reception_path: Optional[Path] = None) -> tuple[bytes, dict]:
    """Generate the report and return (docx_bytes, summary_dict). If
    `reception_path` is given, also fills the VI-3 section from it."""

    df_all = read_excel_data(excel_path)
    df = filter_by_month(df_all, year, month)

    try:
        document = Document(word_path)
    except Exception as e:
        raise TemplateFormatError(f"Could not open the Word template: {e}")

    table_map = locate_tables(document)

    summary = {}
    for rule in HEADING_RULES:
        if rule.category not in selected_categories:
            continue

        cat = rule.category
        cat_df = df[df["BDD"].isin(bdd_values_for_category(cat))]

        table = table_map.get(cat)
        if table is None:
            summary[cat] = {"status": "warning", "msg": "Heading/table not found in document", "rows": 0}
            continue

        if len(cat_df) == 0:
            # Clear all existing data rows so the previous month's data doesn't linger
            existing_data_rows = list(table.rows[1:])
            for _ in range(len(existing_data_rows)):
                remove_row(table, len(table.rows) - 1)
            summary[cat] = {"status": "info", "msg": "No data for this month — old rows cleared", "rows": 0}
            continue

        rows_data = []
        for _, r in cat_df.iterrows():
            rows_data.append({
                "etudes":      r["Data_fmt"],
                "type":        r["Type_fmt"],
                "date":        r["Date_fmt"],
                "source":      r["Source_fmt"],
                "capacite":    r["Capacite_fmt"],
                "realisation": r["Realisation_fmt"],
                "remarque":    r["Remarque_fmt"],
            })

        written = fill_table(table, rows_data, cat)
        summary[cat] = {"status": "ok", "msg": f"{written} row(s) written", "rows": written}

    if reception_path is not None:
        reception_summary = generate_reception_section(document, reception_path)

        t = reception_summary["Terrain"]
        if t["total"] == 0:
            summary["VI-3 Terrain"] = {"status": "info", "msg": "No terrain data in reception file", "rows": 0}
        else:
            parts = [f"{t['matched']}/{t['total']} étude(s) matched"]
            if t["fuzzy_matches"]:
                fuzzy_desc = "; ".join(f"{fm['etude']} — {fm['note']}" for fm in t["fuzzy_matches"])
                parts.append(f"{len(t['fuzzy_matches'])} via partial/fuzzy match, please verify ({fuzzy_desc})")
            if t["unmatched"]:
                unmatched_desc = []
                for u in t["unmatched"]:
                    if u["closest_permis"]:
                        unmatched_desc.append(f"{u['etude']} (closest: \"{u['closest_permis']}\", {u['similarity']:.0%} similar)")
                    else:
                        unmatched_desc.append(u["etude"])
                parts.append(f"not found: {', '.join(unmatched_desc)}")

            status = "warning" if (t["unmatched"] or t["fuzzy_matches"]) else "ok"
            summary["VI-3 Terrain"] = {"status": status, "msg": " — ".join(parts), "rows": t["matched"]}

        s = reception_summary["Stack"]
        if s["total"] == 0:
            summary["VI-3 Stack"] = {"status": "info", "msg": "No stack receptions in reception file", "rows": 0}
        else:
            summary["VI-3 Stack"] = {
                "status": "ok",
                "msg": f"{s['added']} new row(s) added, {s['skipped']} already present",
                "rows": s["added"],
            }

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read(), summary


def preview_data(excel_path: Path, year: int, month: int) -> dict:
    """Return a preview of what would be written, without touching the doc."""
    df_all = read_excel_data(excel_path)
    df = filter_by_month(df_all, year, month)

    preview = {}
    for rule in HEADING_RULES:
        cat = rule.category
        cat_df = df[df["BDD"].isin(bdd_values_for_category(cat))]
        rows = []
        for _, r in cat_df.iterrows():
            rows.append({
                "etudes":      r["Data_fmt"],
                "type":        r["Type_fmt"],
                "date":        r["Date_fmt"],
                "source":      r["Source_fmt"],
                "capacite":    r["Capacite_fmt"],
                "realisation": r["Realisation_fmt"],
                "remarque":    r["Remarque_fmt"],
            })
        preview[cat] = {"label": rule.label, "rows": rows}
    return preview

# ─────────────────────────────────────────────────────────────────────────────
# VI-3 — "RÉCEPTION DES DONNÉES NOUVELLEMENT ACQUISES" SECTION
#
# This section is filled from a SECOND, separate Excel file (the
# "rap_Mens_..." monthly reception file), which has a completely different
# layout from the Situation_Package_DDS file used for the VI-1 sections.
# It is not month-filtered here: the reception file already represents a
# single reporting period end-to-end, so everything found in it is used.
#
# Two independent sub-sections are handled:
#   a. "sismique terrain"  -> VI-3-1 "Réception Données sismiques Terrain"
#        For every Étude found in the Excel, the matching row in the Word
#        table (matched on the "Permis" column) gets APPENDED to — never
#        overwritten — in "Données reçues" and "Nombre Sup 3592".
#   b. "sismique stack"    -> VI-3-2/3 "Réception Données sismiques stack"
#        Every reception found in the Excel becomes a brand-new row
#        appended to the table, unless an identical row (same étude, date,
#        centre) already exists there.
# ─────────────────────────────────────────────────────────────────────────────

RECEPTION_DATA_SUFFIX_RE = re.compile(r"\s*(RAW\s*DATA|CLEAN\s*DATA)\s*$", re.IGNORECASE)


def _norm_row(raw: pd.DataFrame, row_idx: int) -> list[str]:
    return [normalize(x) if pd.notna(x) else "" for x in raw.iloc[row_idx].tolist()]


def _find_marker_row(raw: pd.DataFrame, tokens: tuple, start: int = 0, col: int = 0) -> Optional[int]:
    """Find the first row whose given column contains all tokens (normalised)."""
    for i in range(start, len(raw)):
        val = raw.iat[i, col] if col < raw.shape[1] else None
        norm = normalize(val) if pd.notna(val) else ""
        if norm and all(tok in norm for tok in tokens):
            return i
    return None


def _find_header_row(raw: pd.DataFrame, after: int, token: str = "etude", within: int = 12) -> Optional[int]:
    """Find the first row (searching from `after`) that contains a cell equal to `token`."""
    for i in range(after, min(after + within, len(raw))):
        if token in _norm_row(raw, i):
            return i
    return None


def _combined_headers(raw: pd.DataFrame, header_idx: int) -> list[str]:
    """Reception sheets use a two-row header (e.g. 'Nombre du data reçu en' / 'Cassette').
    Combine both rows into one normalised string per column so keyword lookup works
    regardless of which of the two rows actually carries the label."""
    h1 = raw.iloc[header_idx].tolist()
    h2 = raw.iloc[header_idx + 1].tolist() if header_idx + 1 < len(raw) else []
    h2 = list(h2) + [None] * (len(h1) - len(h2))
    combined = []
    for a, b in zip(h1, h2):
        parts = []
        if pd.notna(a):
            parts.append(normalize(a))
        if b is not None and pd.notna(b):
            parts.append(normalize(b))
        combined.append(" ".join(parts))
    return combined


def _find_col(headers: list[str], tokens: tuple) -> Optional[int]:
    for idx, h in enumerate(headers):
        if all(tok in h for tok in tokens):
            return idx
    return None


def _strip_data_suffix(s) -> str:
    return RECEPTION_DATA_SUFFIX_RE.sub("", str(s).strip()).strip()


def read_terrain_groups(excel_path: Path) -> dict:
    """Read 'a. sismique terrain' and return {study_key(etude): {...}} groups."""
    raw = pd.read_excel(excel_path, header=None)

    start_idx = _find_marker_row(raw, ("sismique", "terrain"))
    if start_idx is None:
        return {}

    header_idx = _find_header_row(raw, start_idx)
    if header_idx is None:
        return {}

    headers = _combined_headers(raw, header_idx)
    col_etude    = _find_col(headers, ("etude",))
    col_numero   = _find_col(headers, ("profils",)) or _find_col(headers, ("numero",))
    col_cassette = _find_col(headers, ("cassette",))
    if col_etude is None:
        return {}

    data_start = header_idx + 2
    groups: dict[str, dict] = {}
    order: list[str] = []

    for i in range(data_start, len(raw)):
        col0 = normalize(raw.iat[i, 0]) if pd.notna(raw.iat[i, 0]) else ""
        if "sismique stack" in col0:
            break
        numero_val = raw.iat[i, col_numero] if col_numero is not None else None
        if isinstance(numero_val, str) and "total" in numero_val.lower():
            break

        etude_val = raw.iat[i, col_etude]
        if pd.isna(etude_val) or str(etude_val).strip() == "":
            continue

        key = study_key(etude_val)
        if not key:
            continue
        if key not in groups:
            groups[key] = {"etude_display": str(etude_val).strip(), "profiles": [], "cassettes": []}
            order.append(key)

        if col_numero is not None and pd.notna(numero_val):
            profile = _strip_data_suffix(numero_val)
            if profile and profile not in groups[key]["profiles"]:
                groups[key]["profiles"].append(profile)

        if col_cassette is not None and pd.notna(raw.iat[i, col_cassette]):
            try:
                groups[key]["cassettes"].append(int(float(raw.iat[i, col_cassette])))
            except (ValueError, TypeError):
                pass

    for key, g in groups.items():
        g["profile_append"]  = " ".join(f"+ {p}" for p in g["profiles"])
        g["cassette_append"] = " ".join(f"+ {c}" for c in g["cassettes"])

    return groups


def read_stack_records(excel_path: Path) -> list[dict]:
    """Read 'b- Sismique stack' and return a list of reception records."""
    raw = pd.read_excel(excel_path, header=None)

    start_idx = _find_marker_row(raw, ("sismique", "stack"))
    if start_idx is None:
        return []

    header_idx = _find_header_row(raw, start_idx)
    if header_idx is None:
        return []

    headers = _combined_headers(raw, header_idx)
    col_dir    = _find_col(headers, ("direction",))
    col_etude  = _find_col(headers, ("etude",))
    col_date   = _find_col(headers, ("date",))
    col_cass   = _find_col(headers, ("cass",))
    col_cd     = _find_col(headers, ("cd",))
    col_doc    = _find_col(headers, ("doc",))
    col_usb    = _find_col(headers, ("usb",))
    col_centre = _find_col(headers, ("centre",))
    if col_etude is None:
        return []

    data_start = header_idx + 2
    records = []

    for i in range(data_start, len(raw)):
        etude_val = raw.iat[i, col_etude]
        if pd.isna(etude_val) or str(etude_val).strip() == "":
            continue

        parts = []
        for col, unit in ((col_cass, "Cartouches"), (col_cd, "CD/DVD"), (col_doc, "CD/DVD"), (col_usb, "USB")):
            if col is None:
                continue
            val = raw.iat[i, col]
            if pd.notna(val):
                try:
                    n = int(float(val))
                    if n != 0:
                        parts.append(f"{n} {unit}")
                except (ValueError, TypeError):
                    pass

        records.append({
            "direction": excel_text(raw.iat[i, col_dir]) if col_dir is not None else "",
            "etude":     str(etude_val).strip(),
            "date":      excel_date(raw.iat[i, col_date]) if col_date is not None else "",
            "centre":    excel_text(raw.iat[i, col_centre]) if col_centre is not None else "",
            "donnees":   " + ".join(parts),
        })

    return records


# ─── Word-side: locating the VI-3 tables (multiple tables per heading) ────────

def locate_reception_tables(document: Document) -> dict[str, list[Table]]:
    """Unlike locate_tables() (VI-1), a VI-3 heading can repeat as '(suite)'
    with its own table each time, so every category maps to a LIST of tables."""
    assigned: dict[str, list[Table]] = {"terrain": [], "stack": []}
    pending: Optional[str] = None

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            norm = normalize(block.text)
            if not norm:
                continue
            if "reception" in norm and "sismique" in norm and "terrain" in norm:
                pending = "terrain"
            elif "reception" in norm and "sismique" in norm and "stack" in norm:
                pending = "stack"
        elif isinstance(block, Table):
            if pending is not None:
                assigned[pending].append(block)
                pending = None

    return assigned


def _find_table_col(table: Table, tokens: tuple) -> Optional[int]:
    for idx, cell in enumerate(table.rows[0].cells):
        norm = normalize(cell.text)
        if all(tok in norm for tok in tokens):
            return idx
    return None


# ─────────────────────────────────────────────────────────────────────────────
# HIGHLIGHTING — flags newly-added reception data (VI-3 Terrain / Stack) so
# it's visually obvious in the generated Word doc without touching anything
# that was already there.
#
# This uses Word's actual "Text Highlight Color" run property
# (<w:highlight w:val="yellow"/>) — the same marker-pen effect Word's
# highlighter button applies, and exactly what's already used by hand on
# the example row in Rap_Mensuel-Template.docx. It's a run-level property,
# so it only ever colors the words it's applied to, never the cell
# background — no separate "word vs cell" logic is needed.
# ─────────────────────────────────────────────────────────────────────────────

HIGHLIGHT_COLOR = "yellow"  # valid w:highlight values: yellow, green, cyan, magenta, blue, red, ...


def _highlight_run(run, color: str = HIGHLIGHT_COLOR):
    """Apply Word's highlighter to a single run's text."""
    rpr = run._r.get_or_add_rPr()
    for existing in rpr.findall(qn("w:highlight")):
        rpr.remove(existing)
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), color)
    rpr.append(hl)


def _highlight_cell_words(cell: _Cell, color: str = HIGHLIGHT_COLOR):
    """Apply Word's highlighter to every run of text inside a cell — used
    for brand-new Stack rows, where the whole row is newly-written data."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                _highlight_run(run, color)


def _append_to_cell(cell: _Cell, addition: str):
    """Append text to a cell's existing content on a new line, preserving
    formatting, without touching anything already in the cell. Idempotent:
    running twice with the same addition will not duplicate it. The
    appended text is placed in its own run and highlighted in yellow so the
    newly-added info stands out from what was already in the cell."""
    if not addition:
        return
    existing_text = cell.text
    if addition.strip() and addition.strip() in existing_text:
        return  # already applied (e.g. report generated twice)

    last_para = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()
    ref_run = None
    for p in cell.paragraphs:
        if p.runs:
            ref_run = p.runs[-1]

    # Always a *new* run for the addition (never extend an existing run),
    # so shading applies only to the newly-added text, not prior content.
    run = last_para.add_run()
    if ref_run is not None:
        ref_rpr = ref_run._r.find(qn("w:rPr"))
        if ref_rpr is not None:
            run._r.insert(0, copy.deepcopy(ref_rpr))

    r = run._r
    if existing_text.strip():
        r.append(OxmlElement("w:br"))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = addition
    r.append(t)

    _highlight_run(run, HIGHLIGHT_COLOR)


def fill_terrain_tables(tables: list[Table], groups: dict) -> dict:
    """Append profile/cassette data into matching rows, keyed on 'Permis'.

    Matching is deliberately conservative — this writes real reception data
    into a specific row, so a wrong match is worse than a missed one:

      1. Exact match on study_key() — handles case, accents, dash/space/
         underscore differences, and stray whitespace. This is the normal
         case and covers the vast majority of études.
      2. Word-order match on study_key_sorted() — handles a trailing
         descriptor written in a different order, e.g. '2025-HMD-Sud-3D'
         vs '2025-HMD-3D SUD'. Still an exact match on the same set of
         tokens, just order-independent, so it's as safe as tier 1.
      3. Containment match — handles a Permis cell with extra annotation
         text around the real code, e.g. '2025-ZAR-3D (Zeml)'. Only applied
         when the shorter of the two keys is at least MIN_CONTAINMENT_LEN
         characters, so a short fragment can't coincidentally match.

    Generic character-similarity ("fuzzy") matching is intentionally NOT
    auto-applied: Étude/Permis codes in this report all share the same
    'YYYY-XXX-ND' shape, so two completely unrelated codes routinely score
    40-60% similar just from the shared digits and dashes — auto-applying
    at that kind of threshold silently writes one étude's data into another
    étude's row. Instead, every étude that doesn't find a safe match is
    reported as unmatched, together with the closest candidate found (and
    its similarity score) purely as a hint for manual review — nothing is
    written for it.
    """
    matched: set[str] = set()
    match_notes: dict[str, str] = {}

    candidates = []
    for table in tables:
        col_permis  = _find_table_col(table, ("permis",))
        col_donnees = _find_table_col(table, ("donnees", "recu"))
        col_sup     = _find_table_col(table, ("sup",))
        if col_permis is None:
            continue
        for row in table.rows[1:]:
            if col_permis >= len(row.cells):
                continue
            permis_text = row.cells[col_permis].text
            key = study_key(permis_text)
            if not key:
                continue
            candidates.append({
                "row": row, "key": key, "sorted_key": study_key_sorted(permis_text), "text": permis_text,
                "col_donnees": col_donnees, "col_sup": col_sup,
            })

    def _apply(cand, g):
        if cand["col_donnees"] is not None and cand["col_donnees"] < len(cand["row"].cells):
            _append_to_cell(cand["row"].cells[cand["col_donnees"]], g["profile_append"])
        if cand["col_sup"] is not None and cand["col_sup"] < len(cand["row"].cells):
            _append_to_cell(cand["row"].cells[cand["col_sup"]], g["cassette_append"])

    # Tier 1 — exact key match (can hit several rows per group, e.g. the
    # same study listed once per team/équipe — both legitimately get the
    # new reception data)
    for gkey, g in groups.items():
        for cand in candidates:
            if cand["key"] == gkey:
                _apply(cand, g)
                matched.add(gkey)

    # Tier 2 — word-order match: same tokens as the group key, just in a
    # different order (e.g. a 'Sud'/'3D' suffix written either way round).
    for gkey, g in groups.items():
        if gkey in matched:
            continue
        g_sorted = study_key_sorted(g["etude_display"])
        if not g_sorted:
            continue
        for cand in candidates:
            if cand["sorted_key"] == g_sorted:
                _apply(cand, g)
                matched.add(gkey)
                match_notes[gkey] = f'word-order match with "{cand["text"].strip()}"'

    # Tier 3 — containment (extra annotation text around the code), gated
    # on both sides so a short fragment can't coincidentally "contain" an
    # unrelated code.
    for gkey, g in groups.items():
        if gkey in matched or len(gkey) < MIN_CONTAINMENT_LEN:
            continue
        for cand in candidates:
            if len(cand["key"]) < MIN_CONTAINMENT_LEN:
                continue
            if gkey in cand["key"] or cand["key"] in gkey:
                _apply(cand, g)
                matched.add(gkey)
                match_notes[gkey] = f'partial match with "{cand["text"].strip()}"'

    unmatched = []
    for k, g in groups.items():
        if k in matched:
            continue
        best_text, best_ratio = None, 0.0
        for cand in candidates:
            r = study_key_similarity(k, cand["key"])
            if r > best_ratio:
                best_text, best_ratio = cand["text"], r
        unmatched.append({
            "etude": g["etude_display"],
            "closest_permis": best_text.strip() if best_text else None,
            "similarity": round(best_ratio, 2) if best_text else None,
        })

    fuzzy_matches = [{"etude": g["etude_display"], "note": match_notes[k]} for k, g in groups.items() if k in match_notes]

    return {
        "matched": len(matched),
        "unmatched": unmatched,
        "total": len(groups),
        "fuzzy_matches": fuzzy_matches,
    }


def fill_stack_tables(tables: list[Table], records: list[dict]) -> dict:
    """Append a new row per reception, skipping ones already present."""
    if not tables:
        return {"added": 0, "skipped": 0, "total": len(records)}

    existing_signatures = set()
    col_maps = []
    for table in tables:
        col_map = {
            "direction": _find_table_col(table, ("direction",)),
            "etude":     _find_table_col(table, ("projet",)),
            "centre":    _find_table_col(table, ("centre",)),
            "date":      _find_table_col(table, ("date",)),
            "donnees":   _find_table_col(table, ("donnees", "recu")),
        }
        col_maps.append(col_map)
        for row in table.rows[1:]:
            def cell_text(field, use_study_key=False):
                idx = col_map.get(field)
                if idx is None or idx >= len(row.cells):
                    return ""
                raw = row.cells[idx].text
                return study_key(raw) if use_study_key else normalize(raw)
            existing_signatures.add((cell_text("etude", use_study_key=True), cell_text("date"), cell_text("centre")))

    target_table = tables[-1]
    target_col_map = col_maps[-1]
    format_template_row_idx = len(target_table.rows) - 1 if len(target_table.rows) > 1 else 0

    added = 0
    skipped = 0
    for rec in records:
        sig = (study_key(rec["etude"]), normalize(rec["date"]), normalize(rec["centre"]))
        if sig in existing_signatures:
            skipped += 1
            continue

        ref_row = target_table.rows[format_template_row_idx]
        new_row = clone_row(target_table, format_template_row_idx)
        for field, col_idx in target_col_map.items():
            if col_idx is not None and col_idx < len(new_row.cells):
                set_cell_text(new_row.cells[col_idx], rec.get(field, ""), reference_cell=ref_row.cells[col_idx])

        for cell in new_row.cells:
            _highlight_cell_words(cell, HIGHLIGHT_COLOR)

        existing_signatures.add(sig)
        added += 1

    return {"added": added, "skipped": skipped, "total": len(records)}


def generate_reception_section(document: Document, reception_excel_path: Path) -> dict:
    """Fill the VI-3 section in-place on `document`. Returns a summary dict."""
    terrain_groups  = read_terrain_groups(reception_excel_path)
    stack_records   = read_stack_records(reception_excel_path)
    table_map       = locate_reception_tables(document)

    summary = {}
    if terrain_groups:
        summary["Terrain"] = fill_terrain_tables(table_map.get("terrain", []), terrain_groups)
    else:
        summary["Terrain"] = {"matched": 0, "unmatched": [], "total": 0}

    if stack_records:
        summary["Stack"] = fill_stack_tables(table_map.get("stack", []), stack_records)
    else:
        summary["Stack"] = {"added": 0, "skipped": 0, "total": 0}

    return summary


def preview_reception_data(reception_excel_path: Path) -> dict:
    """Preview of what the reception file contains, without touching any document."""
    terrain_groups = read_terrain_groups(reception_excel_path)
    stack_records  = read_stack_records(reception_excel_path)
    return {
        "terrain": [
            {
                "etude": g["etude_display"],
                "profile_append": g["profile_append"],
                "cassette_append": g["cassette_append"],
            }
            for g in terrain_groups.values()
        ],
        "stack": stack_records,
    }


# ─────────────────────────────────────────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────────────────────────────────────────

@app.route("/reports/favicon.ico")
def favicon():
    """Silence browser favicon requests — without this every tab open logs an
    [ERROR] 404, burying real errors in noise."""
    return "", 204


@app.errorhandler(AppError)
def handle_app_error(err: AppError):
    logger.warning("AppError: %s", err.message)
    return jsonify({"error": err.message}), err.status_code


@app.errorhandler(Exception)
def handle_unexpected_error(err: Exception):
    if isinstance(err, AppError):
        return handle_app_error(err)
    logger.exception("Unhandled error")
    return jsonify({"error": f"Unexpected error: {err}"}), 500


def _save_upload(file_storage, field_label: str) -> tuple[str, Path]:
    uid = str(uuid.uuid4())
    ext = Path(file_storage.filename).suffix
    save_path = UPLOAD_DIR / f"{uid}{ext}"
    file_storage.save(save_path)
    logger.info("Uploaded %s: '%s' -> %s", field_label, file_storage.filename, save_path.name)
    return uid, save_path


def _get_uploaded_file(file_id: str, label: str) -> Path:
    if not file_id:
        raise FileNotFoundInStoreError(f"No {label} file was provided.")
    matches = list(UPLOAD_DIR.glob(f"{file_id}.*"))
    if not matches:
        raise FileNotFoundInStoreError(f"{label} file not found on the server — please re-upload it.")
    return matches[0]


@app.route("/reports/")
def index_reports():
    categories = [{"code": r.category, "label": r.label} for r in HEADING_RULES]
    return render_template("index.html", categories=categories)


@app.route("/reports/api/upload", methods=["POST"])
def upload_files():
    result = {}

    if "excel" in request.files:
        f = request.files["excel"]
        if f.filename:
            uid, save_path = _save_upload(f, "Situation Package Excel")
            result["excel_id"]   = uid
            result["excel_name"] = f.filename
            try:
                df = read_excel_data(save_path)
                months = get_available_months(df)
                cats   = sorted(df["BDD"].dropna().unique().tolist())
                result["months"]     = months
                result["categories"] = cats
            except AppError as e:
                result["excel_error"] = e.message
                logger.warning("Excel validation failed for '%s': %s", f.filename, e.message)

    if "word" in request.files:
        f = request.files["word"]
        if f.filename:
            uid, save_path = _save_upload(f, "Word template")
            result["word_id"]   = uid
            result["word_name"] = f.filename
            try:
                document = Document(save_path)
                report = validate_template(document)
                result["template_validation"] = report
                logger.info("Template validation for '%s': overall=%s", f.filename, report["overall_status"])
            except Exception as e:
                result["template_error"] = f"Could not open '{f.filename}' as a Word document: {e}"
                logger.warning("Template open failed for '%s': %s", f.filename, e)

    if "reception" in request.files:
        f = request.files["reception"]
        if f.filename:
            uid, save_path = _save_upload(f, "Reception Data Excel")
            result["reception_id"]   = uid
            result["reception_name"] = f.filename
            try:
                terrain_groups = read_terrain_groups(save_path)
                stack_records  = read_stack_records(save_path)
                result["reception_terrain_count"] = len(terrain_groups)
                result["reception_stack_count"]   = len(stack_records)
            except Exception as e:
                result["reception_error"] = str(e)
                logger.warning("Reception Excel validation failed for '%s': %s", f.filename, e)

    return jsonify(result)


@app.route("/reports/api/validate_template", methods=["POST"])
def validate_template_route():
    """Re-run the template validation on demand (e.g. after the user swaps
    the Word file without re-uploading everything else)."""
    data = request.json or {}
    word_path = _get_uploaded_file(data.get("word_id"), "Word template")
    try:
        document = Document(word_path)
    except Exception as e:
        raise TemplateFormatError(f"Could not open the Word template: {e}")
    return jsonify(validate_template(document))


@app.route("/reports/api/preview", methods=["POST"])
def preview():
    data = request.json or {}
    excel_path = _get_uploaded_file(data.get("excel_id"), "Excel")
    try:
        year  = int(data.get("year"))
        month = int(data.get("month"))
    except (TypeError, ValueError):
        raise AppError("A valid year and month must be selected.")
    return jsonify(preview_data(excel_path, year, month))


@app.route("/reports/api/preview_reception", methods=["POST"])
def preview_reception():
    data = request.json or {}
    reception_path = _get_uploaded_file(data.get("reception_id"), "Reception Data Excel")
    return jsonify(preview_reception_data(reception_path))


@app.route("/reports/api/generate", methods=["POST"])
def generate():
    data = request.json or {}
    excel_id = data.get("excel_id")
    word_id  = data.get("word_id")
    try:
        year  = int(data.get("year"))
        month = int(data.get("month"))
    except (TypeError, ValueError):
        raise AppError("A valid year and month must be selected.")
    if not (1 <= month <= 12):
        raise AppError(f"Invalid month: {month}.")

    selected_categories = data.get("categories", [r.category for r in HEADING_RULES])
    reception_id   = data.get("reception_id")
    fill_reception = bool(data.get("fill_reception")) and bool(reception_id)

    excel_path = _get_uploaded_file(excel_id, "Situation Package Excel")
    word_path  = _get_uploaded_file(word_id, "Word template")
    reception_path = _get_uploaded_file(reception_id, "Reception Data Excel") if fill_reception else None

    logger.info(
        "Generate requested: excel=%s word=%s year=%s month=%s categories=%s fill_reception=%s",
        excel_path.name, word_path.name, year, month, selected_categories, fill_reception,
    )

    month_label = datetime(year, month, 1).strftime("%B_%Y")
    docx_bytes, summary = generate_report(
        excel_path, word_path, year, month, selected_categories, reception_path=reception_path
    )

    output_id   = str(uuid.uuid4())
    output_path = OUTPUT_DIR / f"{output_id}.docx"
    output_path.write_bytes(docx_bytes)

    logger.info("Generated report %s (%s) — summary: %s", output_id, month_label, summary)

    return jsonify({
        "output_id": output_id,
        "filename":  f"Situation_Package_{month_label}.docx",
        "summary":   summary,
    })


@app.route("/reports/api/download/<output_id>")
def download(output_id):
    filename    = request.args.get("filename", "report.docx")
    output_path = OUTPUT_DIR / f"{output_id}.docx"
    if not output_path.exists():
        raise FileNotFoundInStoreError("This generated report is no longer available — please generate it again.")
    logger.info("Downloaded report %s as '%s'", output_id, filename)
    return send_file(
        output_path,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )



# Combined main
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
