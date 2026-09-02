#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Sonatrach Situation Package E&P — Report Generator Web App
Flask backend with full report generation logic
"""

from __future__ import annotations

import copy
import io
import logging
import logging.handlers
import os
import re
import unicodedata
import uuid
from dataclasses import dataclass
from datetime import datetime, timedelta
from difflib import SequenceMatcher
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from flask import Flask, jsonify, render_template, request, send_file

app = Flask(__name__)
app.secret_key = os.urandom(24)

UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")
LOG_DIR    = Path("logs")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)
LOG_DIR.mkdir(exist_ok=True)

# ─────────────────────────────────────────────────────────────────────────────
# LOGGING — rotating file log of every upload/validate/generate/download, plus
# console output. No file contents are logged, only filenames/ids/summaries,
# so this is safe to keep around for troubleshooting without becoming a data
# store of its own.
# ─────────────────────────────────────────────────────────────────────────────

logger = logging.getLogger("sonatrach_report")
logger.setLevel(logging.INFO)
if not logger.handlers:
    _file_handler = logging.handlers.RotatingFileHandler(
        LOG_DIR / "app.log", maxBytes=2_000_000, backupCount=5, encoding="utf-8"
    )
    _console_handler = logging.StreamHandler()
    _fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
    _file_handler.setFormatter(_fmt)
    _console_handler.setFormatter(_fmt)
    logger.addHandler(_file_handler)
    logger.addHandler(_console_handler)


# ─────────────────────────────────────────────────────────────────────────────
# ERROR TYPES — used to return specific, actionable messages to the UI
# instead of a generic "Internal Server Error" for anything that goes wrong.
# ─────────────────────────────────────────────────────────────────────────────

class AppError(Exception):
    """Base class for errors we can explain clearly to the user."""
    status_code = 400

    def __init__(self, message: str, status_code: int | None = None):
        super().__init__(message)
        self.message = message
        if status_code is not None:
            self.status_code = status_code


class FileNotFoundInStoreError(AppError):
    status_code = 404


class ExcelFormatError(AppError):
    status_code = 422


class TemplateFormatError(AppError):
    status_code = 422

# ─────────────────────────────────────────────────────────────────────────────
# HEADING RULES
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class HeadingRule:
    category: str
    label: str
    must_contain: tuple
    is_main_section: bool = False


HEADING_RULES: list[HeadingRule] = [
    HeadingRule("PED",        "VI-1- Package de données E&P",           ("vi-1-", "package", "donnees"), is_main_section=True),
    HeadingRule("EDS",        "VI-1-2 Package_EDS",                     ("vi-1-2", "eds")),
    HeadingRule("Asset",      "VI-1-3 Package_Asset",                   ("vi-1-3", "asset")),
    HeadingRule("BDD",        "VI-1-4 Package BDD",                     ("vi-1-4", "bdd")),
    HeadingRule("Traitement", "VI-1-5 Package Traitement/Retraitement", ("vi-1-5",)),
    HeadingRule("SINOPEC",    "VI-1-6 Package_SINOPEC",                 ("vi-1-6", "sinopec")),
    HeadingRule("ENI",        "VI-1-7 Package ENI",                     ("vi-1-7", "eni")),
    HeadingRule("ZPEC",       "VI-1-8 Package ZPEC",                    ("vi-1-8", "zpec")),
    HeadingRule("Total",      "VI-1-9 Package Total",                   ("vi-1-9",)),
    HeadingRule("SH & ZANGAZ&FILADA", "VI-1-10 SH & ZANGAZ&FILADA",     ("vi-1-10",)),
]

CATEGORY_LABELS = {r.category: r.label for r in HEADING_RULES}

# For most categories, the Excel "BDD" column value is just the category
# name itself (e.g. "PED" rows -> PED table). "Asset" is the one exception:
# its table is filled from FOUR separate BDD values, not a single "Asset"
# row. Every other category keeps matching on its own name exactly as before.
CATEGORY_BDD_VALUES: dict[str, list[str]] = {
    "Asset": ["ASCentre", "ASOuest", "ASEst", "ASNord"],
    "Traitement": ["DP/DMG"],
}


def bdd_values_for_category(category: str) -> list[str]:
    return CATEGORY_BDD_VALUES.get(category, [category])

COLUMN_KEYWORDS: dict[str, tuple] = {
    "etudes":      ("etude",),
    "type":        ("type",),
    "date":        ("date",),
    "source":      ("source",),
    "capacite":    ("capacite",),
    "realisation": ("realisation",),
    "remarque":    ("remarque",),
}

EXPECTED_COLUMNS = ["BDD", "Data", "Type_de_donnees", "Date", "Source", "Capacite", "Realisation", "Remarque"]

FIELD_TO_COLUMN = {
    "etudes":      "Data_fmt",
    "type":        "Type_fmt",
    "date":        "Date_fmt",
    "source":      "Source_fmt",
    "capacite":    "Capacite_fmt",
    "realisation": "Realisation_fmt",
    "remarque":    "Remarque_fmt",
}

# ─────────────────────────────────────────────────────────────────────────────
# NORMALISATION HELPERS
# ─────────────────────────────────────────────────────────────────────────────

_DASH_VARIANTS = "\u2010\u2011\u2012\u2013\u2014\u2015\u2212"  # various unicode dash characters


def normalize(text: str) -> str:
    if text is None:
        return ""
    text = unicodedata.normalize("NFKD", str(text))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower()
    for dash in _DASH_VARIANTS:
        text = text.replace(dash, "-")
    text = re.sub(r"[_\u00a0]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def heading_matches(norm_text: str, rule: HeadingRule) -> bool:
    if not norm_text:
        return False
    if rule.is_main_section:
        if not re.match(r"^vi-1-\s", norm_text) and not re.match(r"^vi-1-[^0-9]", norm_text):
            return False
        if re.match(r"^vi-1-\d", norm_text):
            return False
    return all(token in norm_text for token in rule.must_contain)


def study_key(text) -> str:
    """A robust matching key for Étude / Permis codes (e.g. '2025-DJA-2D').
    Strips ALL separators, parenthetical annotations, and casing so that
    formatting differences that don't change the actual code — hyphen vs
    space vs underscore, extra whitespace, a trailing '(Zeml)' style note —
    never cause a genuine match to be missed. Used only for Étude/Permis
    matching, never for heading detection (which needs '-' preserved)."""
    if text is None:
        return ""
    t = normalize(text)
    t = re.sub(r"\([^)]*\)", " ", t)   # drop parenthetical annotations
    t = re.sub(r"[^a-z0-9]", "", t)    # keep only letters/digits
    return t


def study_key_sorted(text) -> str:
    """Order-invariant sibling of study_key(): same normalization, but the
    individual alphanumeric tokens are sorted before joining. This exists
    for one specific, safe purpose — a trailing descriptor written in a
    different word order, e.g. '2025-HMD-Sud-3D' vs '2025-HMD-3D SUD',
    should still count as the same étude. Since it compares the exact same
    multiset of tokens (not a character-similarity score), it can't drift
    into the kind of coincidental match study_key_similarity() warns about:
    two genuinely different codes never share every token, only their
    order. study_key() itself is left untouched — containment matching
    (tier 2) still needs the original left-to-right order preserved so an
    annotation appended after the real code (e.g. '(Zeml)') keeps working."""
    if text is None:
        return ""
    t = normalize(text)
    t = re.sub(r"\([^)]*\)", " ", t)
    tokens = re.findall(r"[a-z0-9]+", t)
    return "".join(sorted(tokens))


def study_key_similarity(a: str, b: str) -> float:
    """0..1 similarity between two study_key() outputs. IMPORTANT: this is
    used only to surface a "closest candidate" hint for a human to review —
    it is NEVER used to automatically decide a match. Étude/Permis codes in
    this dataset (e.g. '2025-DJA-2D', '2026-EAGLE-2D', '2025-YOUCEF-3D') all
    share the same 'YYYY-XXX-ND' shape, so two genuinely different codes
    routinely score 40-60% similar on generic character overlap alone —
    auto-applying at that kind of threshold silently writes real reception
    data (profiles, cassette counts) into the wrong étude's row. See
    MIN_CONTAINMENT_LEN below for the (safe) leniency this script actually
    applies automatically."""
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()


# Minimum length of a study_key() for it to be eligible for the "containment"
# match tier (tier 2) below — e.g. matching '2025ZAR3D' inside
# '2025zar3dzeml'. Below this length, a substring match is too likely to be
# coincidental (e.g. a bare "3d" or "2d") to trust automatically.
MIN_CONTAINMENT_LEN = 6


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT WALKING
# ─────────────────────────────────────────────────────────────────────────────

def iter_block_items(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def locate_tables(document: Document) -> dict[str, Table]:
    assigned: dict[str, Table] = {}
    pending: Optional[HeadingRule] = None

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            norm = normalize(block.text)
            if not norm:
                continue
            for rule in HEADING_RULES:
                if rule.category in assigned:
                    continue
                if heading_matches(norm, rule):
                    pending = rule
                    break
        elif isinstance(block, Table):
            if pending is not None and pending.category not in assigned:
                assigned[pending.category] = block
                pending = None

    return assigned


def locate_table_columns(table: Table) -> dict[str, int]:
    header_cells = table.rows[0].cells
    mapping: dict[str, int] = {}
    for idx, cell in enumerate(header_cells):
        norm = normalize(cell.text)
        for field_name, keywords in COLUMN_KEYWORDS.items():
            if field_name in mapping:
                continue
            if any(kw in norm for kw in keywords):
                mapping[field_name] = idx
    return mapping


# ─────────────────────────────────────────────────────────────────────────────
# TEMPLATE VALIDATION — pre-flight check run right after a Word template is
# uploaded, so a broken/edited template surfaces a clear diagnostic instead
# of a silent "table not found" (or worse, a wrong table) at generation time.
# ─────────────────────────────────────────────────────────────────────────────

def _vi1_heading_found_map(document: Document) -> dict[str, bool]:
    found = {r.category: False for r in HEADING_RULES}
    for block in iter_block_items(document):
        if not isinstance(block, Paragraph):
            continue
        norm = normalize(block.text)
        if not norm:
            continue
        for rule in HEADING_RULES:
            if not found[rule.category] and heading_matches(norm, rule):
                found[rule.category] = True
    return found


def _vi3_heading_found(document: Document) -> tuple[bool, bool]:
    terrain_found = False
    stack_found = False
    for block in iter_block_items(document):
        if not isinstance(block, Paragraph):
            continue
        norm = normalize(block.text)
        if not norm:
            continue
        if "reception" in norm and "sismique" in norm and "terrain" in norm:
            terrain_found = True
        elif "reception" in norm and "sismique" in norm and "stack" in norm:
            stack_found = True
    return terrain_found, stack_found


def _table_missing_columns(table: Table, expected: dict[str, tuple]) -> list[str]:
    missing = []
    for field, tokens in expected.items():
        if _find_table_col(table, tokens) is None:
            missing.append(field)
    return missing


VI3_TERRAIN_EXPECTED_COLUMNS = {
    "permis":          ("permis",),
    "donnees_recues":  ("donnees", "recu"),
    "nombre_sup_3592": ("sup",),
}

VI3_STACK_EXPECTED_COLUMNS = {
    "direction":      ("direction",),
    "projet":         ("projet",),
    "centre":         ("centre",),
    "date_reception": ("date",),
    "donnees_recues": ("donnees", "recu"),
}


def validate_template(document: Document) -> dict:
    """Run a full structural check of the Word template: every VI-1 heading,
    every VI-3 heading, whether a table follows each, and whether the
    expected columns are present in each table found. Never raises — always
    returns a report, so it's safe to call speculatively right after upload."""

    vi1_heading_found = _vi1_heading_found_map(document)
    vi1_table_map = locate_tables(document)

    vi1_report: dict[str, dict] = {}
    for rule in HEADING_RULES:
        cat = rule.category
        table = vi1_table_map.get(cat)
        entry = {
            "label": rule.label,
            "heading_found": vi1_heading_found[cat],
            "table_found": table is not None,
            "missing_columns": None,
            "data_rows": None,
        }
        if table is not None:
            col_map = locate_table_columns(table)
            entry["missing_columns"] = [f for f in COLUMN_KEYWORDS if f not in col_map]
            entry["data_rows"] = len(table.rows) - 1

        if not entry["heading_found"]:
            entry["status"] = "missing"
        elif not entry["table_found"]:
            entry["status"] = "error"
        elif entry["missing_columns"]:
            entry["status"] = "warning"
        else:
            entry["status"] = "ok"
        vi1_report[cat] = entry

    terrain_heading_found, stack_heading_found = _vi3_heading_found(document)
    reception_tables = locate_reception_tables(document)

    terrain_tables = reception_tables.get("terrain", [])
    terrain_table_reports = [
        {"missing_columns": _table_missing_columns(t, VI3_TERRAIN_EXPECTED_COLUMNS), "data_rows": len(t.rows) - 1}
        for t in terrain_tables
    ]
    vi3_terrain = {
        "heading_found": terrain_heading_found,
        "tables_found":  len(terrain_tables),
        "tables":        terrain_table_reports,
    }
    if not terrain_heading_found:
        vi3_terrain["status"] = "missing"
    elif not terrain_tables:
        vi3_terrain["status"] = "error"
    elif any(t["missing_columns"] for t in terrain_table_reports):
        vi3_terrain["status"] = "warning"
    else:
        vi3_terrain["status"] = "ok"

    stack_tables = reception_tables.get("stack", [])
    stack_table_reports = [
        {"missing_columns": _table_missing_columns(t, VI3_STACK_EXPECTED_COLUMNS), "data_rows": len(t.rows) - 1}
        for t in stack_tables
    ]
    vi3_stack = {
        "heading_found": stack_heading_found,
        "tables_found":  len(stack_tables),
        "tables":        stack_table_reports,
    }
    if not stack_heading_found:
        vi3_stack["status"] = "missing"
    elif not stack_tables:
        vi3_stack["status"] = "error"
    elif any(t["missing_columns"] for t in stack_table_reports):
        vi3_stack["status"] = "warning"
    else:
        vi3_stack["status"] = "ok"

    all_statuses = [e["status"] for e in vi1_report.values()] + [vi3_terrain["status"], vi3_stack["status"]]
    if "error" in all_statuses:
        overall = "error"
    elif "warning" in all_statuses:
        overall = "warning"
    else:
        overall = "ok"

    return {
        "vi1":            vi1_report,
        "vi3_terrain":    vi3_terrain,
        "vi3_stack":      vi3_stack,
        "overall_status": overall,
    }


# ─────────────────────────────────────────────────────────────────────────────
# EXCEL READING
# ─────────────────────────────────────────────────────────────────────────────

def excel_date(val) -> str:
    if pd.isna(val) or val == "":
        return ""
    if isinstance(val, datetime):
        return val.strftime("%d/%m/%Y")
    if isinstance(val, (int, float)):
        try:
            return (datetime(1899, 12, 30) + timedelta(days=int(val))).strftime("%d/%m/%Y")
        except Exception:
            return str(val)
    if isinstance(val, str):
        for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%m/%d/%Y"):
            try:
                return datetime.strptime(val.strip(), fmt).strftime("%d/%m/%Y")
            except ValueError:
                pass
        return val.strip()
    return str(val)


def excel_number(val) -> str:
    if pd.isna(val) or val == "":
        return ""
    try:
        f = float(val)
        return str(int(f)) if f == int(f) else str(round(f, 4))
    except (ValueError, TypeError):
        return str(val).strip()


def excel_multiline(val) -> str:
    if pd.isna(val) or val == "":
        return ""
    s = str(val).strip()
    s = re.sub(r"\s*,\s*", "\n", s)
    return s


def excel_text(val) -> str:
    if pd.isna(val) or val == "":
        return ""
    return str(val).strip()


def read_excel_data(excel_path: Path) -> pd.DataFrame:
    try:
        raw = pd.read_excel(excel_path, header=None)
    except Exception as e:
        raise ExcelFormatError(
            f"Could not open '{excel_path.name}' as an Excel file. "
            f"Make sure it's a valid .xlsx/.xls file and isn't corrupted or password-protected. ({e})"
        )

    if raw.empty:
        raise ExcelFormatError(f"'{excel_path.name}' appears to be empty.")

    header_row_idx = None
    for i in range(min(len(raw), 25)):
        first_cell = str(raw.iat[i, 0]).strip().lower() if pd.notna(raw.iat[i, 0]) else ""
        if first_cell == "bdd":
            header_row_idx = i
            break

    if header_row_idx is None:
        raise ExcelFormatError(
            "Could not find the header row (a row starting with 'BDD') in the Excel file. "
            "Check that this is the 'Situation Package' export and that its layout hasn't changed."
        )

    df = raw.iloc[header_row_idx + 1:].reset_index(drop=True)
    n_cols = len(EXPECTED_COLUMNS)
    if df.shape[1] < n_cols:
        raise ExcelFormatError(
            f"Expected at least {n_cols} columns starting from the 'BDD' header row, "
            f"but only found {df.shape[1]}. The Excel layout may have changed."
        )
    df = df.iloc[:, :n_cols]
    df.columns = EXPECTED_COLUMNS

    df = df.dropna(subset=["BDD"])
    df["BDD"] = df["BDD"].astype(str).str.strip()
    df = df[df["BDD"] != "nan"]

    if df.empty:
        raise ExcelFormatError(
            "No usable data rows were found under the 'BDD' header row. "
            "Check that the Excel file has been filled in."
        )

    df["Data_fmt"]        = df["Data"].apply(excel_multiline)
    df["Type_fmt"]        = df["Type_de_donnees"].apply(excel_text)
    df["Date_fmt"]        = df["Date"].apply(excel_date)
    df["Source_fmt"]      = df["Source"].apply(excel_text)
    df["Capacite_fmt"]    = df["Capacite"].apply(excel_number)
    df["Realisation_fmt"] = df["Realisation"].apply(excel_date)
    df["Remarque_fmt"]    = df["Remarque"].apply(excel_text)

    # Store raw date for month filtering
    df["_raw_date"] = pd.to_datetime(df["Date"], errors="coerce")

    return df


def get_available_months(df: pd.DataFrame) -> list[dict]:
    """Return list of {year, month, label} dicts sorted chronologically."""
    months_seen = set()
    for _, row in df.iterrows():
        d = row["_raw_date"]
        if pd.notna(d):
            months_seen.add((d.year, d.month))
    result = []
    for y, m in sorted(months_seen):
        label = datetime(y, m, 1).strftime("%B %Y")
        result.append({"year": y, "month": m, "label": label})
    return result


def filter_by_month(df: pd.DataFrame, year: int, month: int) -> pd.DataFrame:
    mask = df["_raw_date"].apply(
        lambda d: pd.notna(d) and d.year == year and d.month == month
    )
    return df[mask].copy()

# ─────────────────────────────────────────────────────────────────────────────
# CELL / ROW WRITING (format-preserving)
# ─────────────────────────────────────────────────────────────────────────────

def _clear_run_content(run):
    r = run._r
    for child in list(r):
        if child.tag in (qn("w:t"), qn("w:br"), qn("w:tab")):
            r.remove(child)


def _write_lines_into_run(run, lines: list[str]):
    _clear_run_content(run)
    r = run._r
    for i, line in enumerate(lines):
        if i > 0:
            r.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r.append(t)


def _ensure_run_with_format(paragraph, reference_run=None):
    if paragraph.runs:
        for extra in paragraph.runs[1:]:
            extra._r.getparent().remove(extra._r)
        return paragraph.runs[0]
    new_run = paragraph.add_run()
    if reference_run is not None:
        ref_rpr = reference_run._r.find(qn("w:rPr"))
        if ref_rpr is not None:
            new_run._r.insert(0, copy.deepcopy(ref_rpr))
    return new_run


def set_cell_text(cell: _Cell, text: str, reference_cell: Optional[_Cell] = None):
    paragraphs = cell.paragraphs
    for extra_para in paragraphs[1:]:
        extra_para._p.getparent().remove(extra_para._p)
    target_para = cell.paragraphs[0]

    reference_run = None
    if reference_cell is not None and reference_cell.paragraphs and reference_cell.paragraphs[0].runs:
        reference_run = reference_cell.paragraphs[0].runs[0]

    run = _ensure_run_with_format(target_para, reference_run)
    lines = text.split("\n") if text else [""]
    _write_lines_into_run(run, lines)


def clone_row(table: Table, template_row_index: int):
    template_tr = table.rows[template_row_index]._tr
    new_tr = copy.deepcopy(template_tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def remove_row(table: Table, row_index: int):
    tr = table.rows[row_index]._tr
    table._tbl.remove(tr)


def fill_table(table: Table, rows_data: list[dict], category: str) -> int:
    col_map = locate_table_columns(table)
    if not col_map:
        return 0

    header_row = table.rows[0]
    existing_data_rows = list(table.rows[1:])
    needed = len(rows_data)

    # Keep one data row as a formatting template before any trimming/clearing
    # Use the first existing data row (not the header) so we inherit correct
    # font/color/size (e.g. Calibri black) rather than the header style (bold white Arial).
    format_template_row_idx = 1 if len(table.rows) > 1 else 0

    if needed > len(existing_data_rows):
        while len(existing_data_rows) < needed:
            new_row = clone_row(table, format_template_row_idx)
            existing_data_rows.append(new_row)
    elif needed < len(existing_data_rows):
        surplus = len(existing_data_rows) - needed
        for _ in range(surplus):
            remove_row(table, len(table.rows) - 1)
        existing_data_rows = existing_data_rows[:needed]

    # Use the first data row as the formatting reference for new/empty cells.
    # This ensures cloned rows inherit Calibri/Arial black (data style), not
    # the bold white Arial of the header row.
    ref_row = existing_data_rows[0] if existing_data_rows else header_row

    written = 0
    for row_obj, record in zip(existing_data_rows, rows_data):
        for field_name, col_idx in col_map.items():
            if col_idx < len(row_obj.cells):
                cell = row_obj.cells[col_idx]
                ref_cell = ref_row.cells[col_idx]
                value = record.get(field_name, "")
                set_cell_text(cell, value, reference_cell=ref_cell)
        written += 1

    return written

# ─────────────────────────────────────────────────────────────────────────────
# MAIN GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_report(excel_path: Path, word_path: Path, year: int, month: int,
                    selected_categories: list[str],
                    reception_path: Optional[Path] = None) -> tuple[bytes, dict]:
    """Generate the report and return (docx_bytes, summary_dict). If
    `reception_path` is given, also fills the VI-3 section from it."""

    df_all = read_excel_data(excel_path)
    df = filter_by_month(df_all, year, month)

    try:
        document = Document(word_path)
    except Exception as e:
        raise TemplateFormatError(f"Could not open the Word template: {e}")

    table_map = locate_tables(document)

    summary = {}
    for rule in HEADING_RULES:
        if rule.category not in selected_categories:
            continue

        cat = rule.category
        cat_df = df[df["BDD"].isin(bdd_values_for_category(cat))]

        table = table_map.get(cat)
        if table is None:
            summary[cat] = {"status": "warning", "msg": "Heading/table not found in document", "rows": 0}
            continue

        if len(cat_df) == 0:
            # Clear all existing data rows so the previous month's data doesn't linger
            existing_data_rows = list(table.rows[1:])
            for _ in range(len(existing_data_rows)):
                remove_row(table, len(table.rows) - 1)
            summary[cat] = {"status": "info", "msg": "No data for this month — old rows cleared", "rows": 0}
            continue

        rows_data = []
        for _, r in cat_df.iterrows():
            rows_data.append({
                "etudes":      r["Data_fmt"],
                "type":        r["Type_fmt"],
                "date":        r["Date_fmt"],
                "source":      r["Source_fmt"],
                "capacite":    r["Capacite_fmt"],
                "realisation": r["Realisation_fmt"],
                "remarque":    r["Remarque_fmt"],
            })

        written = fill_table(table, rows_data, cat)
        summary[cat] = {"status": "ok", "msg": f"{written} row(s) written", "rows": written}

    if reception_path is not None:
        reception_summary = generate_reception_section(document, reception_path)

        t = reception_summary["Terrain"]
        if t["total"] == 0:
            summary["VI-3 Terrain"] = {"status": "info", "msg": "No terrain data in reception file", "rows": 0}
        else:
            parts = [f"{t['matched']}/{t['total']} étude(s) matched"]
            if t["fuzzy_matches"]:
                fuzzy_desc = "; ".join(f"{fm['etude']} — {fm['note']}" for fm in t["fuzzy_matches"])
                parts.append(f"{len(t['fuzzy_matches'])} via partial/fuzzy match, please verify ({fuzzy_desc})")
            if t["unmatched"]:
                unmatched_desc = []
                for u in t["unmatched"]:
                    if u["closest_permis"]:
                        unmatched_desc.append(f"{u['etude']} (closest: \"{u['closest_permis']}\", {u['similarity']:.0%} similar)")
                    else:
                        unmatched_desc.append(u["etude"])
                parts.append(f"not found: {', '.join(unmatched_desc)}")

            status = "warning" if (t["unmatched"] or t["fuzzy_matches"]) else "ok"
            summary["VI-3 Terrain"] = {"status": status, "msg": " — ".join(parts), "rows": t["matched"]}

        s = reception_summary["Stack"]
        if s["total"] == 0:
            summary["VI-3 Stack"] = {"status": "info", "msg": "No stack receptions in reception file", "rows": 0}
        else:
            summary["VI-3 Stack"] = {
                "status": "ok",
                "msg": f"{s['added']} new row(s) added, {s['skipped']} already present",
                "rows": s["added"],
            }

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read(), summary


def preview_data(excel_path: Path, year: int, month: int) -> dict:
    """Return a preview of what would be written, without touching the doc."""
    df_all = read_excel_data(excel_path)
    df = filter_by_month(df_all, year, month)

    preview = {}
    for rule in HEADING_RULES:
        cat = rule.category
        cat_df = df[df["BDD"].isin(bdd_values_for_category(cat))]
        rows = []
        for _, r in cat_df.iterrows():
            rows.append({
                "etudes":      r["Data_fmt"],
                "type":        r["Type_fmt"],
                "date":        r["Date_fmt"],
                "source":      r["Source_fmt"],
                "capacite":    r["Capacite_fmt"],
                "realisation": r["Realisation_fmt"],
                "remarque":    r["Remarque_fmt"],
            })
        preview[cat] = {"label": rule.label, "rows": rows}
    return preview

# ─────────────────────────────────────────────────────────────────────────────
# VI-3 — "RÉCEPTION DES DONNÉES NOUVELLEMENT ACQUISES" SECTION
#
# This section is filled from a SECOND, separate Excel file (the
# "rap_Mens_..." monthly reception file), which has a completely different
# layout from the Situation_Package_DDS file used for the VI-1 sections.
# It is not month-filtered here: the reception file already represents a
# single reporting period end-to-end, so everything found in it is used.
#
# Two independent sub-sections are handled:
#   a. "sismique terrain"  -> VI-3-1 "Réception Données sismiques Terrain"
#        For every Étude found in the Excel, the matching row in the Word
#        table (matched on the "Permis" column) gets APPENDED to — never
#        overwritten — in "Données reçues" and "Nombre Sup 3592".
#   b. "sismique stack"    -> VI-3-2/3 "Réception Données sismiques stack"
#        Every reception found in the Excel becomes a brand-new row
#        appended to the table, unless an identical row (same étude, date,
#        centre) already exists there.
# ─────────────────────────────────────────────────────────────────────────────

RECEPTION_DATA_SUFFIX_RE = re.compile(r"\s*(RAW\s*DATA|CLEAN\s*DATA)\s*$", re.IGNORECASE)


def _norm_row(raw: pd.DataFrame, row_idx: int) -> list[str]:
    return [normalize(x) if pd.notna(x) else "" for x in raw.iloc[row_idx].tolist()]


def _find_marker_row(raw: pd.DataFrame, tokens: tuple, start: int = 0, col: int = 0) -> Optional[int]:
    """Find the first row whose given column contains all tokens (normalised)."""
    for i in range(start, len(raw)):
        val = raw.iat[i, col] if col < raw.shape[1] else None
        norm = normalize(val) if pd.notna(val) else ""
        if norm and all(tok in norm for tok in tokens):
            return i
    return None


def _find_header_row(raw: pd.DataFrame, after: int, token: str = "etude", within: int = 12) -> Optional[int]:
    """Find the first row (searching from `after`) that contains a cell equal to `token`."""
    for i in range(after, min(after + within, len(raw))):
        if token in _norm_row(raw, i):
            return i
    return None


def _combined_headers(raw: pd.DataFrame, header_idx: int) -> list[str]:
    """Reception sheets use a two-row header (e.g. 'Nombre du data reçu en' / 'Cassette').
    Combine both rows into one normalised string per column so keyword lookup works
    regardless of which of the two rows actually carries the label."""
    h1 = raw.iloc[header_idx].tolist()
    h2 = raw.iloc[header_idx + 1].tolist() if header_idx + 1 < len(raw) else []
    h2 = list(h2) + [None] * (len(h1) - len(h2))
    combined = []
    for a, b in zip(h1, h2):
        parts = []
        if pd.notna(a):
            parts.append(normalize(a))
        if b is not None and pd.notna(b):
            parts.append(normalize(b))
        combined.append(" ".join(parts))
    return combined


def _find_col(headers: list[str], tokens: tuple) -> Optional[int]:
    for idx, h in enumerate(headers):
        if all(tok in h for tok in tokens):
            return idx
    return None


def _strip_data_suffix(s) -> str:
    return RECEPTION_DATA_SUFFIX_RE.sub("", str(s).strip()).strip()


def read_terrain_groups(excel_path: Path) -> dict:
    """Read 'a. sismique terrain' and return {study_key(etude): {...}} groups."""
    raw = pd.read_excel(excel_path, header=None)

    start_idx = _find_marker_row(raw, ("sismique", "terrain"))
    if start_idx is None:
        return {}

    header_idx = _find_header_row(raw, start_idx)
    if header_idx is None:
        return {}

    headers = _combined_headers(raw, header_idx)
    col_etude    = _find_col(headers, ("etude",))
    col_numero   = _find_col(headers, ("profils",)) or _find_col(headers, ("numero",))
    col_cassette = _find_col(headers, ("cassette",))
    if col_etude is None:
        return {}

    data_start = header_idx + 2
    groups: dict[str, dict] = {}
    order: list[str] = []

    for i in range(data_start, len(raw)):
        col0 = normalize(raw.iat[i, 0]) if pd.notna(raw.iat[i, 0]) else ""
        if "sismique stack" in col0:
            break
        numero_val = raw.iat[i, col_numero] if col_numero is not None else None
        if isinstance(numero_val, str) and "total" in numero_val.lower():
            break

        etude_val = raw.iat[i, col_etude]
        if pd.isna(etude_val) or str(etude_val).strip() == "":
            continue

        key = study_key(etude_val)
        if not key:
            continue
        if key not in groups:
            groups[key] = {"etude_display": str(etude_val).strip(), "profiles": [], "cassettes": []}
            order.append(key)

        if col_numero is not None and pd.notna(numero_val):
            profile = _strip_data_suffix(numero_val)
            if profile and profile not in groups[key]["profiles"]:
                groups[key]["profiles"].append(profile)

        if col_cassette is not None and pd.notna(raw.iat[i, col_cassette]):
            try:
                groups[key]["cassettes"].append(int(float(raw.iat[i, col_cassette])))
            except (ValueError, TypeError):
                pass

    for key, g in groups.items():
        g["profile_append"]  = " ".join(f"+ {p}" for p in g["profiles"])
        g["cassette_append"] = " ".join(f"+ {c}" for c in g["cassettes"])

    return groups


def read_stack_records(excel_path: Path) -> list[dict]:
    """Read 'b- Sismique stack' and return a list of reception records."""
    raw = pd.read_excel(excel_path, header=None)

    start_idx = _find_marker_row(raw, ("sismique", "stack"))
    if start_idx is None:
        return []

    header_idx = _find_header_row(raw, start_idx)
    if header_idx is None:
        return []

    headers = _combined_headers(raw, header_idx)
    col_dir    = _find_col(headers, ("direction",))
    col_etude  = _find_col(headers, ("etude",))
    col_date   = _find_col(headers, ("date",))
    col_cass   = _find_col(headers, ("cass",))
    col_cd     = _find_col(headers, ("cd",))
    col_doc    = _find_col(headers, ("doc",))
    col_usb    = _find_col(headers, ("usb",))
    col_centre = _find_col(headers, ("centre",))
    if col_etude is None:
        return []

    data_start = header_idx + 2
    records = []

    for i in range(data_start, len(raw)):
        etude_val = raw.iat[i, col_etude]
        if pd.isna(etude_val) or str(etude_val).strip() == "":
            continue

        parts = []
        for col, unit in ((col_cass, "Cartouches"), (col_cd, "CD/DVD"), (col_doc, "CD/DVD"), (col_usb, "USB")):
            if col is None:
                continue
            val = raw.iat[i, col]
            if pd.notna(val):
                try:
                    n = int(float(val))
                    if n != 0:
                        parts.append(f"{n} {unit}")
                except (ValueError, TypeError):
                    pass

        records.append({
            "direction": excel_text(raw.iat[i, col_dir]) if col_dir is not None else "",
            "etude":     str(etude_val).strip(),
            "date":      excel_date(raw.iat[i, col_date]) if col_date is not None else "",
            "centre":    excel_text(raw.iat[i, col_centre]) if col_centre is not None else "",
            "donnees":   " + ".join(parts),
        })

    return records


# ─── Word-side: locating the VI-3 tables (multiple tables per heading) ────────

def locate_reception_tables(document: Document) -> dict[str, list[Table]]:
    """Unlike locate_tables() (VI-1), a VI-3 heading can repeat as '(suite)'
    with its own table each time, so every category maps to a LIST of tables."""
    assigned: dict[str, list[Table]] = {"terrain": [], "stack": []}
    pending: Optional[str] = None

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            norm = normalize(block.text)
            if not norm:
                continue
            if "reception" in norm and "sismique" in norm and "terrain" in norm:
                pending = "terrain"
            elif "reception" in norm and "sismique" in norm and "stack" in norm:
                pending = "stack"
        elif isinstance(block, Table):
            if pending is not None:
                assigned[pending].append(block)
                pending = None

    return assigned


def _find_table_col(table: Table, tokens: tuple) -> Optional[int]:
    for idx, cell in enumerate(table.rows[0].cells):
        norm = normalize(cell.text)
        if all(tok in norm for tok in tokens):
            return idx
    return None


# ─────────────────────────────────────────────────────────────────────────────
# HIGHLIGHTING — flags newly-added reception data (VI-3 Terrain / Stack) so
# it's visually obvious in the generated Word doc without touching anything
# that was already there.
#
# This uses Word's actual "Text Highlight Color" run property
# (<w:highlight w:val="yellow"/>) — the same marker-pen effect Word's
# highlighter button applies, and exactly what's already used by hand on
# the example row in Rap_Mensuel-Template.docx. It's a run-level property,
# so it only ever colors the words it's applied to, never the cell
# background — no separate "word vs cell" logic is needed.
# ─────────────────────────────────────────────────────────────────────────────

HIGHLIGHT_COLOR = "yellow"  # valid w:highlight values: yellow, green, cyan, magenta, blue, red, ...


def _highlight_run(run, color: str = HIGHLIGHT_COLOR):
    """Apply Word's highlighter to a single run's text."""
    rpr = run._r.get_or_add_rPr()
    for existing in rpr.findall(qn("w:highlight")):
        rpr.remove(existing)
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), color)
    rpr.append(hl)


def _highlight_cell_words(cell: _Cell, color: str = HIGHLIGHT_COLOR):
    """Apply Word's highlighter to every run of text inside a cell — used
    for brand-new Stack rows, where the whole row is newly-written data."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                _highlight_run(run, color)


def _append_to_cell(cell: _Cell, addition: str):
    """Append text to a cell's existing content on a new line, preserving
    formatting, without touching anything already in the cell. Idempotent:
    running twice with the same addition will not duplicate it. The
    appended text is placed in its own run and highlighted in yellow so the
    newly-added info stands out from what was already in the cell."""
    if not addition:
        return
    existing_text = cell.text
    if addition.strip() and addition.strip() in existing_text:
        return  # already applied (e.g. report generated twice)

    last_para = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()
    ref_run = None
    for p in cell.paragraphs:
        if p.runs:
            ref_run = p.runs[-1]

    # Always a *new* run for the addition (never extend an existing run),
    # so shading applies only to the newly-added text, not prior content.
    run = last_para.add_run()
    if ref_run is not None:
        ref_rpr = ref_run._r.find(qn("w:rPr"))
        if ref_rpr is not None:
            run._r.insert(0, copy.deepcopy(ref_rpr))

    r = run._r
    if existing_text.strip():
        r.append(OxmlElement("w:br"))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = addition
    r.append(t)

    _highlight_run(run, HIGHLIGHT_COLOR)


def fill_terrain_tables(tables: list[Table], groups: dict) -> dict:
    """Append profile/cassette data into matching rows, keyed on 'Permis'.

    Matching is deliberately conservative — this writes real reception data
    into a specific row, so a wrong match is worse than a missed one:

      1. Exact match on study_key() — handles case, accents, dash/space/
         underscore differences, and stray whitespace. This is the normal
         case and covers the vast majority of études.
      2. Word-order match on study_key_sorted() — handles a trailing
         descriptor written in a different order, e.g. '2025-HMD-Sud-3D'
         vs '2025-HMD-3D SUD'. Still an exact match on the same set of
         tokens, just order-independent, so it's as safe as tier 1.
      3. Containment match — handles a Permis cell with extra annotation
         text around the real code, e.g. '2025-ZAR-3D (Zeml)'. Only applied
         when the shorter of the two keys is at least MIN_CONTAINMENT_LEN
         characters, so a short fragment can't coincidentally match.

    Generic character-similarity ("fuzzy") matching is intentionally NOT
    auto-applied: Étude/Permis codes in this report all share the same
    'YYYY-XXX-ND' shape, so two completely unrelated codes routinely score
    40-60% similar just from the shared digits and dashes — auto-applying
    at that kind of threshold silently writes one étude's data into another
    étude's row. Instead, every étude that doesn't find a safe match is
    reported as unmatched, together with the closest candidate found (and
    its similarity score) purely as a hint for manual review — nothing is
    written for it.
    """
    matched: set[str] = set()
    match_notes: dict[str, str] = {}

    candidates = []
    for table in tables:
        col_permis  = _find_table_col(table, ("permis",))
        col_donnees = _find_table_col(table, ("donnees", "recu"))
        col_sup     = _find_table_col(table, ("sup",))
        if col_permis is None:
            continue
        for row in table.rows[1:]:
            if col_permis >= len(row.cells):
                continue
            permis_text = row.cells[col_permis].text
            key = study_key(permis_text)
            if not key:
                continue
            candidates.append({
                "row": row, "key": key, "sorted_key": study_key_sorted(permis_text), "text": permis_text,
                "col_donnees": col_donnees, "col_sup": col_sup,
            })

    def _apply(cand, g):
        if cand["col_donnees"] is not None and cand["col_donnees"] < len(cand["row"].cells):
            _append_to_cell(cand["row"].cells[cand["col_donnees"]], g["profile_append"])
        if cand["col_sup"] is not None and cand["col_sup"] < len(cand["row"].cells):
            _append_to_cell(cand["row"].cells[cand["col_sup"]], g["cassette_append"])

    # Tier 1 — exact key match (can hit several rows per group, e.g. the
    # same study listed once per team/équipe — both legitimately get the
    # new reception data)
    for gkey, g in groups.items():
        for cand in candidates:
            if cand["key"] == gkey:
                _apply(cand, g)
                matched.add(gkey)

    # Tier 2 — word-order match: same tokens as the group key, just in a
    # different order (e.g. a 'Sud'/'3D' suffix written either way round).
    for gkey, g in groups.items():
        if gkey in matched:
            continue
        g_sorted = study_key_sorted(g["etude_display"])
        if not g_sorted:
            continue
        for cand in candidates:
            if cand["sorted_key"] == g_sorted:
                _apply(cand, g)
                matched.add(gkey)
                match_notes[gkey] = f'word-order match with "{cand["text"].strip()}"'

    # Tier 3 — containment (extra annotation text around the code), gated
    # on both sides so a short fragment can't coincidentally "contain" an
    # unrelated code.
    for gkey, g in groups.items():
        if gkey in matched or len(gkey) < MIN_CONTAINMENT_LEN:
            continue
        for cand in candidates:
            if len(cand["key"]) < MIN_CONTAINMENT_LEN:
                continue
            if gkey in cand["key"] or cand["key"] in gkey:
                _apply(cand, g)
                matched.add(gkey)
                match_notes[gkey] = f'partial match with "{cand["text"].strip()}"'

    unmatched = []
    for k, g in groups.items():
        if k in matched:
            continue
        best_text, best_ratio = None, 0.0
        for cand in candidates:
            r = study_key_similarity(k, cand["key"])
            if r > best_ratio:
                best_text, best_ratio = cand["text"], r
        unmatched.append({
            "etude": g["etude_display"],
            "closest_permis": best_text.strip() if best_text else None,
            "similarity": round(best_ratio, 2) if best_text else None,
        })

    fuzzy_matches = [{"etude": g["etude_display"], "note": match_notes[k]} for k, g in groups.items() if k in match_notes]

    return {
        "matched": len(matched),
        "unmatched": unmatched,
        "total": len(groups),
        "fuzzy_matches": fuzzy_matches,
    }


def fill_stack_tables(tables: list[Table], records: list[dict]) -> dict:
    """Append a new row per reception, skipping ones already present."""
    if not tables:
        return {"added": 0, "skipped": 0, "total": len(records)}

    existing_signatures = set()
    col_maps = []
    for table in tables:
        col_map = {
            "direction": _find_table_col(table, ("direction",)),
            "etude":     _find_table_col(table, ("projet",)),
            "centre":    _find_table_col(table, ("centre",)),
            "date":      _find_table_col(table, ("date",)),
            "donnees":   _find_table_col(table, ("donnees", "recu")),
        }
        col_maps.append(col_map)
        for row in table.rows[1:]:
            def cell_text(field, use_study_key=False):
                idx = col_map.get(field)
                if idx is None or idx >= len(row.cells):
                    return ""
                raw = row.cells[idx].text
                return study_key(raw) if use_study_key else normalize(raw)
            existing_signatures.add((cell_text("etude", use_study_key=True), cell_text("date"), cell_text("centre")))

    target_table = tables[-1]
    target_col_map = col_maps[-1]
    format_template_row_idx = len(target_table.rows) - 1 if len(target_table.rows) > 1 else 0

    added = 0
    skipped = 0
    for rec in records:
        sig = (study_key(rec["etude"]), normalize(rec["date"]), normalize(rec["centre"]))
        if sig in existing_signatures:
            skipped += 1
            continue

        ref_row = target_table.rows[format_template_row_idx]
        new_row = clone_row(target_table, format_template_row_idx)
        for field, col_idx in target_col_map.items():
            if col_idx is not None and col_idx < len(new_row.cells):
                set_cell_text(new_row.cells[col_idx], rec.get(field, ""), reference_cell=ref_row.cells[col_idx])

        for cell in new_row.cells:
            _highlight_cell_words(cell, HIGHLIGHT_COLOR)

        existing_signatures.add(sig)
        added += 1

    return {"added": added, "skipped": skipped, "total": len(records)}


def generate_reception_section(document: Document, reception_excel_path: Path) -> dict:
    """Fill the VI-3 section in-place on `document`. Returns a summary dict."""
    terrain_groups  = read_terrain_groups(reception_excel_path)
    stack_records   = read_stack_records(reception_excel_path)
    table_map       = locate_reception_tables(document)

    summary = {}
    if terrain_groups:
        summary["Terrain"] = fill_terrain_tables(table_map.get("terrain", []), terrain_groups)
    else:
        summary["Terrain"] = {"matched": 0, "unmatched": [], "total": 0}

    if stack_records:
        summary["Stack"] = fill_stack_tables(table_map.get("stack", []), stack_records)
    else:
        summary["Stack"] = {"added": 0, "skipped": 0, "total": 0}

    return summary


def preview_reception_data(reception_excel_path: Path) -> dict:
    """Preview of what the reception file contains, without touching any document."""
    terrain_groups = read_terrain_groups(reception_excel_path)
    stack_records  = read_stack_records(reception_excel_path)
    return {
        "terrain": [
            {
                "etude": g["etude_display"],
                "profile_append": g["profile_append"],
                "cassette_append": g["cassette_append"],
            }
            for g in terrain_groups.values()
        ],
        "stack": stack_records,
    }


# ─────────────────────────────────────────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────────────────────────────────────────

@app.route("/favicon.ico")
def favicon():
    """Silence browser favicon requests — without this every tab open logs an
    [ERROR] 404, burying real errors in noise."""
    return "", 204


@app.errorhandler(AppError)
def handle_app_error(err: AppError):
    logger.warning("AppError: %s", err.message)
    return jsonify({"error": err.message}), err.status_code


@app.errorhandler(Exception)
def handle_unexpected_error(err: Exception):
    if isinstance(err, AppError):
        return handle_app_error(err)
    logger.exception("Unhandled error")
    return jsonify({"error": f"Unexpected error: {err}"}), 500


def _save_upload(file_storage, field_label: str) -> tuple[str, Path]:
    uid = str(uuid.uuid4())
    ext = Path(file_storage.filename).suffix
    save_path = UPLOAD_DIR / f"{uid}{ext}"
    file_storage.save(save_path)
    logger.info("Uploaded %s: '%s' -> %s", field_label, file_storage.filename, save_path.name)
    return uid, save_path


def _get_uploaded_file(file_id: str, label: str) -> Path:
    if not file_id:
        raise FileNotFoundInStoreError(f"No {label} file was provided.")
    matches = list(UPLOAD_DIR.glob(f"{file_id}.*"))
    if not matches:
        raise FileNotFoundInStoreError(f"{label} file not found on the server — please re-upload it.")
    return matches[0]


@app.route("/")
def index():
    categories = [{"code": r.category, "label": r.label} for r in HEADING_RULES]
    return render_template("index.html", categories=categories)


@app.route("/api/upload", methods=["POST"])
def upload_files():
    result = {}

    if "excel" in request.files:
        f = request.files["excel"]
        if f.filename:
            uid, save_path = _save_upload(f, "Situation Package Excel")
            result["excel_id"]   = uid
            result["excel_name"] = f.filename
            try:
                df = read_excel_data(save_path)
                months = get_available_months(df)
                cats   = sorted(df["BDD"].dropna().unique().tolist())
                result["months"]     = months
                result["categories"] = cats
            except AppError as e:
                result["excel_error"] = e.message
                logger.warning("Excel validation failed for '%s': %s", f.filename, e.message)

    if "word" in request.files:
        f = request.files["word"]
        if f.filename:
            uid, save_path = _save_upload(f, "Word template")
            result["word_id"]   = uid
            result["word_name"] = f.filename
            try:
                document = Document(save_path)
                report = validate_template(document)
                result["template_validation"] = report
                logger.info("Template validation for '%s': overall=%s", f.filename, report["overall_status"])
            except Exception as e:
                result["template_error"] = f"Could not open '{f.filename}' as a Word document: {e}"
                logger.warning("Template open failed for '%s': %s", f.filename, e)

    if "reception" in request.files:
        f = request.files["reception"]
        if f.filename:
            uid, save_path = _save_upload(f, "Reception Data Excel")
            result["reception_id"]   = uid
            result["reception_name"] = f.filename
            try:
                terrain_groups = read_terrain_groups(save_path)
                stack_records  = read_stack_records(save_path)
                result["reception_terrain_count"] = len(terrain_groups)
                result["reception_stack_count"]   = len(stack_records)
            except Exception as e:
                result["reception_error"] = str(e)
                logger.warning("Reception Excel validation failed for '%s': %s", f.filename, e)

    return jsonify(result)


@app.route("/api/validate_template", methods=["POST"])
def validate_template_route():
    """Re-run the template validation on demand (e.g. after the user swaps
    the Word file without re-uploading everything else)."""
    data = request.json or {}
    word_path = _get_uploaded_file(data.get("word_id"), "Word template")
    try:
        document = Document(word_path)
    except Exception as e:
        raise TemplateFormatError(f"Could not open the Word template: {e}")
    return jsonify(validate_template(document))


@app.route("/api/preview", methods=["POST"])
def preview():
    data = request.json or {}
    excel_path = _get_uploaded_file(data.get("excel_id"), "Excel")
    try:
        year  = int(data.get("year"))
        month = int(data.get("month"))
    except (TypeError, ValueError):
        raise AppError("A valid year and month must be selected.")
    return jsonify(preview_data(excel_path, year, month))


@app.route("/api/preview_reception", methods=["POST"])
def preview_reception():
    data = request.json or {}
    reception_path = _get_uploaded_file(data.get("reception_id"), "Reception Data Excel")
    return jsonify(preview_reception_data(reception_path))


@app.route("/api/generate", methods=["POST"])
def generate():
    data = request.json or {}
    excel_id = data.get("excel_id")
    word_id  = data.get("word_id")
    try:
        year  = int(data.get("year"))
        month = int(data.get("month"))
    except (TypeError, ValueError):
        raise AppError("A valid year and month must be selected.")
    if not (1 <= month <= 12):
        raise AppError(f"Invalid month: {month}.")

    selected_categories = data.get("categories", [r.category for r in HEADING_RULES])
    reception_id   = data.get("reception_id")
    fill_reception = bool(data.get("fill_reception")) and bool(reception_id)

    excel_path = _get_uploaded_file(excel_id, "Situation Package Excel")
    word_path  = _get_uploaded_file(word_id, "Word template")
    reception_path = _get_uploaded_file(reception_id, "Reception Data Excel") if fill_reception else None

    logger.info(
        "Generate requested: excel=%s word=%s year=%s month=%s categories=%s fill_reception=%s",
        excel_path.name, word_path.name, year, month, selected_categories, fill_reception,
    )

    month_label = datetime(year, month, 1).strftime("%B_%Y")
    docx_bytes, summary = generate_report(
        excel_path, word_path, year, month, selected_categories, reception_path=reception_path
    )

    output_id   = str(uuid.uuid4())
    output_path = OUTPUT_DIR / f"{output_id}.docx"
    output_path.write_bytes(docx_bytes)

    logger.info("Generated report %s (%s) — summary: %s", output_id, month_label, summary)

    return jsonify({
        "output_id": output_id,
        "filename":  f"Situation_Package_{month_label}.docx",
        "summary":   summary,
    })


@app.route("/api/download/<output_id>")
def download(output_id):
    filename    = request.args.get("filename", "report.docx")
    output_path = OUTPUT_DIR / f"{output_id}.docx"
    if not output_path.exists():
        raise FileNotFoundInStoreError("This generated report is no longer available — please generate it again.")
    logger.info("Downloaded report %s as '%s'", output_id, filename)
    return send_file(
        output_path,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    # host="0.0.0.0" makes the app reachable from other machines on the
    # local network (via this computer's LAN IP, e.g. http://192.168.x.x:5001),
    # not just from this computer itself.
    app.run(debug=True, host="0.0.0.0", port=5001)
